import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 12, "color": "#000000", "val": "single"},
        start={"sz": 12, "val": "single"},
        end={"sz": 12, "val": "single"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_cell_bg_color(cell, color_str):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_str)
    cell._tc.get_or_add_tcPr().append(shading)

input_file = sys.argv[1] if len(sys.argv) > 1 else "FSD_Master_Data_New_Format.docx"
output_file = sys.argv[2] if len(sys.argv) > 2 else input_file

doc = Document(input_file)

# Set Normal font to Calibri
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Iterate through tables and apply styles
for table in doc.tables:
    for row_idx, row in enumerate(table.rows):
        # Apply borders to all cells
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 4, "val": "single", "color": "000000"},
                bottom={"sz": 4, "val": "single", "color": "000000"},
                start={"sz": 4, "val": "single", "color": "000000"},
                end={"sz": 4, "val": "single", "color": "000000"}
            )
            # Change font of each paragraph to Calibri
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
        
        # Header Row Background (Light Green)
        if row_idx == 0:
            for cell in row.cells:
                set_cell_bg_color(cell, "C6E0B4")
                # Make header font bold
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

doc.save(output_file)
print("Successfully styled document")
