"""
make_reference.py
Creates a pandoc reference.docx with:
- Calibri font for all body text and headings
- Tables with borders and light green (#C6EFCE) header rows
Usage: python make_reference.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

# ---------- helpers ----------

def set_font(run_or_para, name='Calibri', size=11, bold=False, color=None):
    """Apply font on a Run. If a paragraph, apply to all runs."""
    if hasattr(run_or_para, 'runs'):
        for r in run_or_para.runs:
            _apply_font(r, name, size, bold, color)
    else:
        _apply_font(run_or_para, name, size, bold, color)

def _apply_font(run, name, size, bold, color):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_cell_bg(cell, hex_color):
    """Set cell background fill colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table, color='4F7A28', sz=4):
    """Add visible borders to all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), str(sz))
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), color)
                tcBorders.append(border)
            tcPr.append(tcBorders)

def style_heading(doc, style_name, size, bold=True):
    """Set Calibri font on a built-in heading style."""
    try:
        style = doc.styles[style_name]
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(size)
        font.bold = bold
        font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # dark blue, standard Word
    except Exception:
        pass

# ---------- build reference doc ----------

doc = Document()

# Default body style
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

# Heading styles
style_heading(doc, 'Heading 1', 16)
style_heading(doc, 'Heading 2', 14)
style_heading(doc, 'Heading 3', 12)
style_heading(doc, 'Heading 4', 11)

# Add a sample table so pandoc picks up the style
doc.add_heading('Reference Document — IDC FSD', 0)
p = doc.add_paragraph('This document sets the default styling for IDC FSD exports.')
p.runs[0].font.name = 'Calibri'
p.runs[0].font.size = Pt(11)

table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'
set_table_borders(table)

headers = ['Column A', 'Column B', 'Column C']
for i, hdr in enumerate(headers):
    cell = table.rows[0].cells[i]
    set_cell_bg(cell, 'C6EFCE')  # light green
    run = cell.paragraphs[0].add_run(hdr)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True

for ri in range(1, 3):
    for ci in range(3):
        cell = table.rows[ri].cells[ci]
        run = cell.paragraphs[0].add_run(f'Row {ri} Col {ci+1}')
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

doc.save('reference.docx')
print('reference.docx created successfully.')
