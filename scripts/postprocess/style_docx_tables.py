"""
style_docx_tables.py
Post-processes the pandoc-generated DOCX to apply:
- Table borders to all tables
- Light green (#C6EFCE) background on first row (header) of each table
- Calibri 11pt font throughout the document
Run: python style_docx_tables.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = 'FSD_SEAL_Project_Identity_v1.0.docx'
OUTPUT = 'FSD_SEAL_Project_Identity_v1.0.docx'
HEADER_BG = 'C6EFCE'   # light green
BORDER_COLOR = '92D050'  # medium green border
FONT_NAME = 'Calibri'
FONT_SIZE = 11

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shd if any
    for old_shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(old_shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color, sz=8):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_run_font(run, name, size, bold=None):
    run.font.name = name
    run.font.size = Pt(size)
    # Propagate to East-Asian / CS fonts via XML
    rPr = run._r.get_or_add_rPr()
    for tag in (qn('w:rFonts'),):
        el = rPr.find(tag)
        if el is None:
            el = OxmlElement('w:rFonts')
            rPr.insert(0, el)
        el.set(qn('w:ascii'), name)
        el.set(qn('w:hAnsi'), name)
        el.set(qn('w:cs'), name)
    if bold is not None:
        run.font.bold = bold

def apply_font_to_para(para, name=FONT_NAME, size=FONT_SIZE):
    for run in para.runs:
        set_run_font(run, name, size)

def process():
    doc = Document(INPUT)

    # 1. Apply Calibri to all Normal paragraphs
    for para in doc.paragraphs:
        apply_font_to_para(para)

    # 2. Style every table
    for table in doc.tables:
        for ri, row in enumerate(table.rows):
            is_header = (ri == 0)
            for cell in row.cells:
                # Border
                set_cell_borders(cell, BORDER_COLOR)
                # Header background
                if is_header:
                    set_cell_bg(cell, HEADER_BG)
                # Font in cell
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_run_font(run, FONT_NAME, FONT_SIZE, bold=is_header if not run.font.bold else True)
                    # If cell has no runs but has text via direct XML, add run
                    if not para.runs and para.text.strip():
                        run = para.add_run(para.text)
                        para.clear()
                        set_run_font(run, FONT_NAME, FONT_SIZE, bold=is_header)
                        para.add_run()

        # Make table use 'Table Grid' style for borders
        table.style = doc.styles['Table Grid']

    doc.save(OUTPUT)
    print(f'Styled DOCX saved to {OUTPUT}')

if __name__ == '__main__':
    process()
