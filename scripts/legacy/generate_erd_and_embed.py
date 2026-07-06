"""
generate_erd_and_embed.py
1. Generate ERD diagram sebagai PNG menggunakan matplotlib (offline, no internet needed)
2. Embed hasilnya ke FSD_New_RM_Sample_v1.8.docx menggantikan placeholder "screenshot belum tersedia"

Run: python generate_erd_and_embed.py
"""

import os
import re
import matplotlib
matplotlib.use('Agg')  # non-interactive backend
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
SCREENSHOTS = os.path.join(SCRIPT_DIR, 'screenshots')
ERD_PNG     = os.path.join(SCREENSHOTS, 'new_rm_sample_erd.png')
DOCX_PATH   = os.path.join(SCRIPT_DIR, 'FSD_New_RM_Sample_v1.8.docx')

# ── Styling constants (same as build script) ───────────────────────────────────
HEADER_BG    = 'FFF2CC'
BORDER_COLOR = 'BF9000'
FONT_NAME    = 'Calibri'
FONT_SIZE    = 11


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 – Generate ERD as PNG using matplotlib
# ══════════════════════════════════════════════════════════════════════════════

def draw_entity(ax, x, y, width, height, name, fields, header_color='#2E74B5', row_color='#F5F5F5'):
    """Draw a single entity box with header and field rows."""
    row_h = 0.38
    total_h = row_h + len(fields) * row_h

    # Header
    hdr = FancyBboxPatch((x, y - row_h), width, row_h,
                          boxstyle="square,pad=0", linewidth=1,
                          edgecolor='#333333', facecolor=header_color)
    ax.add_patch(hdr)
    ax.text(x + width / 2, y - row_h / 2, name,
            ha='center', va='center', fontsize=8.5, fontweight='bold',
            color='white', fontfamily='DejaVu Sans')

    # Field rows
    for i, (fname, ftype, tag) in enumerate(fields):
        ry = y - row_h - (i + 1) * row_h
        fc = '#E8F0FE' if tag in ('PK', 'UK') else '#FFF8E1' if tag == 'FK' else row_color
        rect = FancyBboxPatch((x, ry), width, row_h,
                               boxstyle="square,pad=0", linewidth=0.5,
                               edgecolor='#BBBBBB', facecolor=fc)
        ax.add_patch(rect)
        # Tag badge
        if tag:
            badge_color = '#1565C0' if tag == 'PK' else '#E65100' if tag == 'FK' else '#2E7D32'
            ax.text(x + 0.08, ry + row_h / 2, tag,
                    ha='left', va='center', fontsize=6, fontweight='bold',
                    color=badge_color, fontfamily='DejaVu Sans')
        ax.text(x + 1.0, ry + row_h / 2, fname,
                ha='left', va='center', fontsize=7, color='#222222',
                fontfamily='DejaVu Sans')
        ax.text(x + width - 0.08, ry + row_h / 2, ftype,
                ha='right', va='center', fontsize=6.5, color='#555555',
                style='italic', fontfamily='DejaVu Sans')

    return total_h


def draw_arrow(ax, x1, y1, x2, y2, label=''):
    """Draw a relationship arrow."""
    ax.annotate('',
                xy=(x2, y2), xycoords='data',
                xytext=(x1, y1), textcoords='data',
                arrowprops=dict(arrowstyle='->', color='#555555',
                                lw=1.2, connectionstyle='arc3,rad=0.0'))
    if label:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        ax.text(mx, my + 0.1, label, ha='center', va='bottom',
                fontsize=6, color='#666666', fontfamily='DejaVu Sans',
                bbox=dict(boxstyle='round,pad=0.15', fc='white', ec='none', alpha=0.8))


def generate_erd():
    os.makedirs(SCREENSHOTS, exist_ok=True)

    fig, ax = plt.subplots(figsize=(18, 12))
    ax.set_xlim(0, 18)
    ax.set_ylim(-11, 2)
    ax.axis('off')
    ax.set_facecolor('#FAFAFA')
    fig.patch.set_facecolor('#FAFAFA')

    # Title
    ax.text(9, 1.3, 'New RM Sample Management – Entity Relationship Diagram',
            ha='center', va='center', fontsize=13, fontweight='bold',
            color='#1A237E', fontfamily='DejaVu Sans')
    ax.text(9, 0.7, 'Database Schema v1.8  |  IDC System',
            ha='center', va='center', fontsize=9, color='#555555',
            fontfamily='DejaVu Sans')

    row_h = 0.38

    # ── HEADER (center) ────────────────────────────────────────────────────────
    hdr_fields = [
        ('intSampleId',      'BIGSERIAL', 'PK'),
        ('txtSampleNo',      'VARCHAR(50)', 'UK'),
        ('intTypeDocumentId','BIGINT', 'FK'),
        ('txtMaterialName',  'VARCHAR(200)', ''),
        ('intSupplierId',    'BIGINT', 'FK'),
        ('txtSupplierName',  'VARCHAR(200)', ''),
        ('decPrice',         'DECIMAL(18,4)', ''),
        ('intCurrencyId',    'BIGINT', 'FK'),
        ('intUOMId',         'BIGINT', 'FK'),
        ('decQuantity',      'DECIMAL(18,4)', ''),
        ('intPICId',         'BIGINT', 'FK'),
        ('dtSubmissionDate', 'TIMESTAMP', ''),
        ('txtWorkflowStage', 'VARCHAR(50)', ''),
        ('txtStatus',        'VARCHAR(50)', ''),
        ('bitActive',        'BOOLEAN', ''),
        ('txtCreatedBy',     'VARCHAR(50)', ''),
        ('dtCreatedDate',    'TIMESTAMP', ''),
    ]
    hdr_x, hdr_y, hdr_w = 6.5, 0.2, 5.0
    hdr_h = draw_entity(ax, hdr_x, hdr_y, hdr_w, row_h,
                         'T_RM_SAMPLE_HEADER', hdr_fields,
                         header_color='#1A237E')

    # ── ATTACHMENT (top-left) ──────────────────────────────────────────────────
    att_fields = [
        ('intAttachId',     'BIGSERIAL', 'PK'),
        ('intSampleId',     'BIGINT', 'FK'),
        ('intDocTypeId',    'BIGINT', 'FK'),
        ('txtFileName',     'VARCHAR(255)', ''),
        ('txtFilePath',     'TEXT', ''),
        ('intFileSizeKb',   'INT', ''),
        ('txtFileType',     'VARCHAR(20)', ''),
        ('txtUploadedBy',   'VARCHAR(50)', ''),
        ('dtUploadedDate',  'TIMESTAMP', ''),
        ('bitActive',       'BOOLEAN', ''),
        ('txtCreatedBy',    'VARCHAR(50)', ''),
    ]
    att_x, att_y, att_w = 0.3, 0.2, 4.8
    draw_entity(ax, att_x, att_y, att_w, row_h,
                'T_RM_SAMPLE_ATTACHMENT', att_fields,
                header_color='#E65100')

    # ── PURPOSE (top-right) ────────────────────────────────────────────────────
    pur_fields = [
        ('intPurposeId',         'BIGSERIAL', 'PK'),
        ('intSampleId',          'BIGINT', 'FK'),
        ('intPurposeCategoryId', 'BIGINT', 'FK'),
        ('txtTargetProduct',     'TEXT', ''),
        ('txtExpectedBenefit',   'VARCHAR(1000)', ''),
        ('decUsagePercentage',   'DECIMAL(5,2)', ''),
        ('intReplaceMaterialId', 'BIGINT', 'FK'),
        ('dtAnalysisDeadline',   'DATE', ''),
        ('txtCreatedBy',         'VARCHAR(50)', ''),
    ]
    pur_x, pur_y, pur_w = 12.9, 0.2, 4.8
    draw_entity(ax, pur_x, pur_y, pur_w, row_h,
                'T_RM_SAMPLE_PURPOSE', pur_fields,
                header_color='#17A2B8')

    # ── EVALUATION (bottom-left) ───────────────────────────────────────────────
    eva_fields = [
        ('intEvalId',           'BIGSERIAL', 'PK'),
        ('intSampleId',         'BIGINT', 'FK'),
        ('dtTestDate',          'DATE', ''),
        ('intTestTypeId',       'BIGINT', 'FK'),
        ('txtTestResult',       'VARCHAR(10)', ''),
        ('txtSpecCompliance',   'VARCHAR(10)', ''),
        ('intQualityScore',     'INT', ''),
        ('txtEvaluationNotes',  'TEXT', ''),
        ('txtPotensiEGDEG',     'VARCHAR(500)', ''),
        ('intEvaluatorId',      'BIGINT', 'FK'),
        ('txtCreatedBy',        'VARCHAR(50)', ''),
    ]
    eva_x, eva_y, eva_w = 0.3, -5.8, 4.8
    draw_entity(ax, eva_x, eva_y, eva_w, row_h,
                'T_RM_SAMPLE_EVALUATION', eva_fields,
                header_color='#6C757D')

    # ── DISPOSITION (bottom-right) ────────────────────────────────────────────
    dis_fields = [
        ('intDispId',            'BIGSERIAL', 'PK'),
        ('intSampleId',          'BIGINT', 'FK'),
        ('txtDecision',          'VARCHAR(20)', ''),
        ('dtDecisionDate',       'TIMESTAMP', ''),
        ('intDecisionById',      'BIGINT', 'FK'),
        ('txtReason',            'VARCHAR(1000)', ''),
        ('txtNextAction',        'VARCHAR(500)', ''),
        ('bitApprovalRequired',  'BOOLEAN', ''),
        ('intApproverId',        'BIGINT', 'FK'),
        ('txtApprovalStatus',    'VARCHAR(20)', ''),
        ('dtApprovalDate',       'TIMESTAMP', ''),
        ('txtApprovalNotes',     'TEXT', ''),
        ('txtCreatedBy',         'VARCHAR(50)', ''),
    ]
    dis_x, dis_y, dis_w = 12.9, -5.8, 4.8
    draw_entity(ax, dis_x, dis_y, dis_w, row_h,
                'T_RM_SAMPLE_DISPOSITION', dis_fields,
                header_color='#DC3545')

    # ── RELATIONSHIP ARROWS ────────────────────────────────────────────────────
    # Header -> Attachment
    draw_arrow(ax, hdr_x, hdr_y - row_h * 5,
               att_x + att_w, att_y - row_h * 2, '1 to N')
    # Header -> Purpose
    draw_arrow(ax, hdr_x + hdr_w, hdr_y - row_h * 5,
               pur_x, pur_y - row_h * 2, '1 to 1')
    # Header -> Evaluation
    draw_arrow(ax, hdr_x, hdr_y - hdr_h,
               eva_x + eva_w, eva_y + row_h * 2, '1 to N')
    # Header -> Disposition
    draw_arrow(ax, hdr_x + hdr_w, hdr_y - hdr_h,
               dis_x, dis_y + row_h * 2, '1 to 1')

    # ── LEGEND ────────────────────────────────────────────────────────────────
    legend_items = [
        mpatches.Patch(color='#1A237E', label='Header Table (Main)'),
        mpatches.Patch(color='#E65100', label='Attachment Table'),
        mpatches.Patch(color='#17A2B8', label='Purpose Table'),
        mpatches.Patch(color='#6C757D', label='Evaluation Table'),
        mpatches.Patch(color='#DC3545', label='Disposition Table'),
    ]
    ax.legend(handles=legend_items, loc='lower left',
              fontsize=7.5, framealpha=0.9,
              bbox_to_anchor=(0.01, 0.01),
              ncol=5)

    # ── PK/FK legend ──────────────────────────────────────────────────────────
    ax.text(0.3, -10.7,
            '[PK] = Primary Key    [FK] = Foreign Key    [UK] = Unique Key'
            '    Blue highlight = PK/UK    Orange highlight = FK',
            ha='left', va='center', fontsize=7.5, color='#444444',
            fontfamily='DejaVu Sans')

    plt.tight_layout(pad=0.5)
    plt.savefig(ERD_PNG, dpi=150, bbox_inches='tight',
                facecolor=fig.get_facecolor())
    plt.close()
    print(f'[OK] ERD saved -> {ERD_PNG} ({os.path.getsize(ERD_PNG)/1024:.0f} KB)')


# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 – Embed ERD PNG into existing DOCX
# ══════════════════════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color='BF9000', sz=8):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def embed_erd_into_docx():
    """Find placeholder paragraph for Gambar 7 / ERD and replace with image."""
    if not os.path.exists(ERD_PNG):
        print('[SKIP] ERD PNG not found, skipping embed.')
        return

    doc = Document(DOCX_PATH)
    erd_embedded = False

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # Look for the placeholder text we put in when erd png was missing
        if ('Gambar 7' in text and ('ERD' in text or 'screenshot belum tersedia' in text or 'ERD – Relasi Tabel' in text)):
            print(f'  Found ERD placeholder at paragraph {i}: "{text[:80]}"')
            # Clear existing runs
            for run in para.runs:
                run.text = ''
            # Add image in first run
            run = para.add_run()
            run.add_picture(ERD_PNG, width=Inches(6.0))
            # Add caption in next paragraph
            cap_para = para._element
            cap = doc.add_paragraph()
            cap.add_run('Gambar 7: ERD – Relasi Tabel Modul New RM Sample Management').italic = True
            cap.alignment = 1  # CENTER
            # Move caption right after image paragraph
            para._element.addnext(cap._element)
            erd_embedded = True
            break

    if not erd_embedded:
        # If placeholder not found, append at end of section 4.2
        print('  Placeholder not found — appending ERD after section 4.2 heading search...')
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if '4.2' in text and 'ERD' in text.upper():
                print(f'  Found section 4.2 at paragraph {i}: "{text[:60]}"')
                # Insert image paragraph after this heading
                img_para = doc.add_paragraph()
                run = img_para.add_run()
                run.add_picture(ERD_PNG, width=Inches(6.0))
                img_para.alignment = 1
                para._element.addnext(img_para._element)
                # Caption
                cap = doc.add_paragraph()
                cap.add_run('Gambar 7: ERD – Relasi Tabel Modul New RM Sample Management').italic = True
                cap.alignment = 1
                img_para._element.addnext(cap._element)
                erd_embedded = True
                break

    if erd_embedded:
        doc.save(DOCX_PATH)
        size_kb = os.path.getsize(DOCX_PATH) / 1024
        print(f'[OK] DOCX saved with ERD -> {DOCX_PATH} ({size_kb:.0f} KB)')
    else:
        print('[WARN] Could not find insertion point. Appending to end of document.')
        # Last resort: append at end
        doc.add_heading('Entity Relationship Diagram', level=3)
        img_para = doc.add_paragraph()
        run = img_para.add_run()
        run.add_picture(ERD_PNG, width=Inches(6.0))
        img_para.alignment = 1
        cap = doc.add_paragraph()
        cap.add_run('Gambar 7: ERD – Relasi Tabel Modul New RM Sample Management').italic = True
        cap.alignment = 1
        doc.save(DOCX_PATH)
        size_kb = os.path.getsize(DOCX_PATH) / 1024
        print(f'[OK] DOCX saved (appended ERD) -> {DOCX_PATH} ({size_kb:.0f} KB)')


# ══════════════════════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print('=' * 60)
    print('  Step 1: Generating ERD diagram with matplotlib...')
    print('=' * 60)
    generate_erd()

    print()
    print('=' * 60)
    print('  Step 2: Embedding ERD into DOCX...')
    print('=' * 60)
    embed_erd_into_docx()

    print()
    print('Done!')
