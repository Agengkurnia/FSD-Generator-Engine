"""
build_fsd_v3.py
===============
Pipeline lengkap untuk membangun FSD New RM Sample v3.0 dalam format DOCX:

  Langkah 1 – Update ERD di Markdown (replace section 5.2 dgn ERD lengkap)
  Langkah 2 – Render diagram Mermaid ke PNG via mermaid.ink
  Langkah 3 – Pre-process MD: ganti blok ```mermaid dan path gambar lokal
  Langkah 4 – Jalankan pandoc untuk buat DOCX
  Langkah 5 – Post-process DOCX: border tabel, warna header, font Calibri

Run:
    py build_fsd_v3.py
"""

import base64
import os
import re
import subprocess
import urllib.request
import shutil

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────────
# PATHS
# ─────────────────────────────────────────────
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
MD_SRC      = os.path.join(SCRIPT_DIR, 'FSD_New_RM_Sample_v3.0.md')
MD_TMP      = os.path.join(SCRIPT_DIR, '_tmp_fsd_v3_processed.md')
DOCX_OUT    = os.path.join(SCRIPT_DIR, 'FSD_New_RM_Sample_v3.0.docx')
REF_DOCX    = os.path.join(SCRIPT_DIR, 'reference.docx')
SCREENSHOTS = os.path.join(SCRIPT_DIR, 'screenshots')

# Brain screenshots (captured from actual UI)
BRAIN_DIR   = r'C:\Users\Lenovo\.gemini\antigravity\brain\76591413-77ac-4e43-9749-e54abab6dd70'

# ─────────────────────────────────────────────
# STYLING
# ─────────────────────────────────────────────
HEADER_BG    = 'D9EAD3'   # light green (like reference FSD)
BORDER_COLOR = '000000'   # black
FONT_NAME    = 'Calibri'
FONT_SIZE_BODY  = 11
FONT_SIZE_TABLE = 9

# ─────────────────────────────────────────────
# MERMAID DIAGRAMS
# ─────────────────────────────────────────────
ERD_WORKFLOW_CODE = """
erDiagram
    trRMSample_Header {
        int intId PK
        varchar txtId UK
        varchar txtSampleNo UK
        int intStatusId
        int intCurrentStep
        varchar txtCreatedBy
        bit bitActive
    }
    trDocumentSample_Header {
        int intId PK
        int intSampleHeaderId FK
        date dtSampleDate
        date dtDateReceipt
        varchar txtMaterialName
        varchar txtSupplierName
        varchar txtGMOStatus
        bit bitContainPHO
        varchar txtStatusOrganik
        varchar txtEgDegContent
        int intShelfLifeMonth
    }
    trDocumentSample_Allergen {
        int intId PK
        int intSampleHeaderId FK
        int intDocSampleHeaderId FK
        varchar txtAllergenCode
        bit bitIsChecked
    }
    trDocumentSample_BTP {
        int intId PK
        int intSampleHeaderId FK
        int intDocSampleHeaderId FK
        int intLineNo
        varchar txtBTPName
        varchar txtFunctionName
        decimal decPercentagePPM
    }
    trDocumentSample_Document {
        int intId PK
        int intSampleHeaderId FK
        varchar txtDocumentType
        varchar txtFileName
        varchar txtFilePath
        bit bitActive
    }
    trSamplePurpose_Header {
        int intId PK
        int intSampleHeaderId FK
        varchar txtRMCategoryCode
        varchar txtSamplePurposeType
        varchar txtItemCodeTrial
        varchar txtItemCodeExisting
        varchar txtObjective
    }
    trSamplePurpose_Product {
        int intId PK
        int intSampleHeaderId FK
        int intSamplePurposeHeaderId FK
        int intLineNo
        varchar txtGroupProductTypeName
        varchar txtChildProductTypeName
        varchar txtVariantName
    }
    trRMEvaluation_Header {
        int intId PK
        int intSampleHeaderId FK
        decimal decOverallScore
        varchar txtEvaluationStatus
        varchar txtMaterialComposition
    }
    trRMEvaluation_Detail {
        int intId PK
        int intEvaluationHeaderId FK
        varchar txtTestCode
        varchar txtTestClass
        varchar txtAnalysisResult
    }
    trRMEvaluation_FoodCategory {
        int intId PK
        int intEvaluationHeaderId FK
        varchar txtCategoryCode
        varchar txtCategoryName
    }
    trRMDisposition_Header {
        int intId PK
        int intSampleHeaderId FK
        varchar txtDecision
        datetime dtDecisionDate
        bit bitApprovalRequired
        varchar txtApprovalStatus
    }

    trRMSample_Header ||--o| trDocumentSample_Header : "Step1"
    trRMSample_Header ||--o{ trDocumentSample_Allergen : "Allergen"
    trRMSample_Header ||--o{ trDocumentSample_BTP : "BTP"
    trRMSample_Header ||--o{ trDocumentSample_Document : "Documents"
    trRMSample_Header ||--o| trSamplePurpose_Header : "Step2"
    trSamplePurpose_Header ||--o{ trSamplePurpose_Product : "Products"
    trRMSample_Header ||--o| trRMEvaluation_Header : "Step3"
    trRMEvaluation_Header ||--o{ trRMEvaluation_Detail : "Parameters"
    trRMEvaluation_Header ||--o{ trRMEvaluation_FoodCategory : "FoodCat"
    trRMSample_Header ||--o| trRMDisposition_Header : "Step4"
    trDocumentSample_Header ||--o{ trDocumentSample_Allergen : "parent"
    trDocumentSample_Header ||--o{ trDocumentSample_BTP : "parent"
"""

ERD_MASTER_CODE = """
erDiagram
    mSupplier {
        int intId PK
        varchar txtSupplierId UK
        varchar txtSupplierName
        varchar txtCountryCode
        bit bitActive
    }
    mHalalCategory {
        int intId PK
        varchar txtHalalCategoryCode UK
        varchar txtHalalCategoryName
        bit bitActive
    }
    mHalalBody {
        int intId PK
        varchar txtHalalBodyInstitution
        varchar txtCountry
        bit bitActive
    }
    mAppParam {
        int intId PK
        varchar txtAppParamVariable
        varchar txtAppParamCode
        varchar txtAppParamValue
        bit bitActive
    }
    mPegawai {
        int intId PK
        varchar txtPegawaiCode UK
        varchar txtPegawaiName
        int intDepartmentId
        varchar txtRole
        bit bitActive
    }
    mGroupProductType {
        int intId PK
        varchar txtGroupProductTypeName
        bit bitActive
    }
    mGroupProductTypeDetail {
        int intId PK
        int intGroupProductTypeId FK
        varchar txtProductBrand
        varchar txtProductCategory
        bit bitActive
    }
    mChildProductType {
        int intId PK
        int intGroupProductTypeId FK
        varchar txtChildProductTypeName
        bit bitActive
    }
    mVarian {
        int intId PK
        int intChildProductTypeId FK
        varchar txtVarianName
        bit bitActive
    }
    trDocumentSample_Header {
        int intSupplierId FK
        int intHalalBodyId FK
        varchar txtHalalCategoryCode FK
        int intStorageId FK
        int intPegawaiId FK
    }
    trSamplePurpose_Product {
        int intGroupProductTypeId FK
        int intGroupProductTypeDetailId FK
        int intChildProductTypeId FK
        int intVariantId FK
    }
    trRMDisposition_Header {
        int intApproverId FK
    }

    trDocumentSample_Header }o--|| mSupplier : "FK Supplier"
    trDocumentSample_Header }o--|| mHalalCategory : "FK HalalCat"
    trDocumentSample_Header }o--|| mHalalBody : "FK HalalBody"
    trDocumentSample_Header }o--|| mPegawai : "FK PIC"
    trDocumentSample_Header }o--|| mAppParam : "FK Storage"
    trRMDisposition_Header }o--|| mPegawai : "FK Approver"
    trSamplePurpose_Product }o--|| mGroupProductType : "FK Group"
    trSamplePurpose_Product }o--|| mGroupProductTypeDetail : "FK GroupDetail"
    trSamplePurpose_Product }o--|| mChildProductType : "FK Child"
    trSamplePurpose_Product }o--|| mVarian : "FK Varian"
    mGroupProductType ||--o{ mGroupProductTypeDetail : "has Detail"
    mGroupProductType ||--o{ mChildProductType : "has Child"
    mChildProductType ||--o{ mVarian : "has Varian"
"""

FLOW_MAIN_CODE = """
flowchart TD
    subgraph Row1 ["Tahap 1: Evaluasi"]
        direction LR
        Start(["Mulai"]) --> Step1["STEP 1: DOCUMENT REGISTRATION"] --> Step2["STEP 2: DOCUMENT SAMPLE"] --> Step3["STEP 3: RM EVALUATION"]
    end
    Step3 --> Step4["STEP 4: DISPOSITION"]
    subgraph Row2 ["Tahap 2: Keputusan"]
        direction RL
        Step4 --> Decision{"Keputusan?"}
        Decision --> |"APPROVED"| ApprovalCheck{"Perlu Approval?"}
        Decision --> |"REJECTED"| EndReject(["REJECTED"])
        Decision --> |"ON HOLD"| OnHold["ON HOLD"]
        Decision --> |"CANCELLED"| EndCancel(["CANCELLED"])
    end
    OnHold -.-> |"Kembali"| Step3
    ApprovalCheck --> PendingApproval["Pending Approval"]
    ApprovalCheck --> |"Tidak"| Completed1(["COMPLETED"])
    subgraph Row3 ["Tahap 3: Approval"]
        direction LR
        PendingApproval --> ApproverReview["Approver Review"] --> ApprovalDecision{"Hasil?"}
        ApprovalDecision --> |"APPROVE"| Completed(["COMPLETED"])
    end
    ApprovalDecision -.-> |"REJECT"| Step4
    style Start fill:#8CC63F,color:#fff,stroke:#009B4D
    style Completed fill:#388E3C,color:#fff,stroke:#1B5E20
    style Completed1 fill:#388E3C,color:#fff,stroke:#1B5E20
    style EndReject fill:#D32F2F,color:#fff,stroke:#B71C1C
    style EndCancel fill:#9E9E9E,color:#fff,stroke:#616161
    style Step1 fill:#FFF3CD,stroke:#FFC107
    style Step2 fill:#D1ECF1,stroke:#17A2B8
    style Step3 fill:#E2E3E5,stroke:#6C757D
    style Step4 fill:#F8D7DA,stroke:#DC3545
"""

# Output PNG locations
os.makedirs(SCREENSHOTS, exist_ok=True)
ERD_WORKFLOW_PNG = os.path.join(SCREENSHOTS, 'erd_workflow_v3.png')
ERD_MASTER_PNG   = os.path.join(SCREENSHOTS, 'erd_master_v3.png')
FLOW_MAIN_PNG    = os.path.join(SCREENSHOTS, 'flow_main_v3.png')

# Mapping UI screenshots from brain dir
UI_SCREENSHOT_MAP = {
    'step1_header_wizard':       os.path.join(BRAIN_DIR, 'step1_header_wizard_1775702646454.png'),
    'step1_supplier_material':   os.path.join(BRAIN_DIR, 'step1_supplier_material_expanded_1775702658871.png'),
    'step1_pricing_packaging':   os.path.join(BRAIN_DIR, 'step1_pricing_packaging_expanded_1775702707845.png'),
    'step1_storage_shelflife':   os.path.join(BRAIN_DIR, 'step1_storage_shelflife_expanded_1775702732270.png'),
    'step1_halal_gmo':           os.path.join(BRAIN_DIR, 'step1_halal_gmo_expanded_1775702779578.png'),
    'step1_allergen':            os.path.join(BRAIN_DIR, 'step1_allergen_expanded_1775702814071.png'),
    'step1_btp_content':         os.path.join(BRAIN_DIR, 'step1_btp_content_expanded_1775702840503.png'),
    'step2_full_view':           os.path.join(BRAIN_DIR, 'step2_full_view_1775702898567.png'),
    'step2_lov_popup':           os.path.join(BRAIN_DIR, 'step2_lov_popup_1775702929514.png'),
    'step2_document_upload':     os.path.join(BRAIN_DIR, 'step2_document_upload_expanded_1775703115155.png'),
}

# ═════════════════════════════════════════════
# STEP 1 – Render Mermaid diagrams to PNG
# ═════════════════════════════════════════════
def render_mermaid_png(mermaid_code, output_path, label):
    encoded = base64.urlsafe_b64encode(mermaid_code.strip().encode('utf-8')).decode('ascii')
    url = f'https://mermaid.ink/img/{encoded}?type=png&scale=2'
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = resp.read()
        with open(output_path, 'wb') as f:
            f.write(data)
        print(f'   [{label}] ✓ {len(data):,} bytes → {os.path.basename(output_path)}')
        return True
    except Exception as e:
        print(f'   [{label}] ✗ mermaid.ink gagal: {e}')
        return False


def step1_render_diagrams():
    print('\n[STEP 1] Render diagram Mermaid -> PNG')
    render_mermaid_png(ERD_WORKFLOW_CODE,  ERD_WORKFLOW_PNG,  'ERD-Workflow')
    render_mermaid_png(ERD_MASTER_CODE,    ERD_MASTER_PNG,    'ERD-Master')
    render_mermaid_png(FLOW_MAIN_CODE,     FLOW_MAIN_PNG,     'Flow-Main')


# ═════════════════════════════════════════════
# STEP 2 – Copy UI screenshots to screenshots/
# ═════════════════════════════════════════════
def step2_copy_screenshots():
    print('\n[STEP 2] Menyalin screenshot UI ke folder screenshots/')
    for key, src in UI_SCREENSHOT_MAP.items():
        dst = os.path.join(SCREENSHOTS, f'{key}.png')
        if os.path.exists(src):
            shutil.copy2(src, dst)
            print(f'   ✓ {key}.png')
        else:
            print(f'   ✗ Tidak ditemukan: {src}')


# ═════════════════════════════════════════════
# STEP 3 – Pre-process Markdown
# ═════════════════════════════════════════════
def _rel(path):
    """Return relative path from SCRIPT_DIR using forward slashes."""
    return os.path.relpath(path, SCRIPT_DIR).replace('\\', '/')

def step3_preprocess_markdown():
    print('\n[STEP 3] Pre-processing Markdown...')
    with open(MD_SRC, 'r', encoding='utf-8') as f:
        text = f.read()

    # 1. Replace mermaid code blocks with PNG images
    counter = [0]
    def replace_mermaid_block(m):
        counter[0] += 1
        code = m.group(1).strip()
        # Decide which PNG to use based on diagram type
        if 'erDiagram' in code:
            if 'mSupplier' in code or 'mAppParam' in code:
                png = ERD_MASTER_PNG
                caption = 'ERD – Relasi ke Tabel Master'
            else:
                png = ERD_WORKFLOW_PNG
                caption = 'ERD – Transaksi Workflow'
        elif 'flowchart LR' in code and 'Submit' in code:
            # small inline flow — skip, render as text note
            return '\n> *[Diagram: Logika Submit – lihat keterangan di atas]*\n'
        else:
            png = FLOW_MAIN_PNG
            caption = 'Business Flow Diagram'

        if os.path.exists(png):
            return f'\n![{caption}]({_rel(png)})\n'
        return f'\n*({caption} – diagram tidak tersedia)*\n'

    text = re.sub(r'```mermaid\s*\n(.*?)```', replace_mermaid_block, text, flags=re.DOTALL)

    # 2. Replace absolute screenshot paths with relative paths
    def replace_abs_img(m):
        full_path = m.group(2).replace('\\', '/')
        fname = os.path.basename(full_path)
        # Try to find matching file in screenshots dir
        local_candidates = [
            os.path.join(SCREENSHOTS, fname),
        ]
        # Also match by key name patterns
        for key in UI_SCREENSHOT_MAP:
            if key in fname or any(k in fname for k in ['step1', 'step2']):
                candidate = os.path.join(SCREENSHOTS, f'{key}.png')
                local_candidates.append(candidate)

        for candidate in local_candidates:
            if os.path.exists(candidate):
                return f'![{m.group(1)}]({_rel(candidate)})'

        # Fallback: try the original path if it's accessible
        if os.path.exists(full_path):
            return f'![{m.group(1)}]({full_path})'

        return f'*[Screenshot: {m.group(1)} – tidak tersedia]*'

    text = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', replace_abs_img, text)

    with open(MD_TMP, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f'   ✓ MD sementara: {os.path.basename(MD_TMP)}')


# ═════════════════════════════════════════════
# STEP 4 – Jalankan Pandoc
# ═════════════════════════════════════════════
def step4_run_pandoc():
    print('\n[STEP 4] Menjalankan pandoc...')
    cmd = [
        'pandoc', MD_TMP,
        '-o', DOCX_OUT,
        '--from=markdown+pipe_tables',
        f'--resource-path={SCRIPT_DIR};{SCREENSHOTS}',
        '--toc',
        '--toc-depth=3',
    ]
    if os.path.exists(REF_DOCX):
        cmd += [f'--reference-doc={REF_DOCX}']
        print(f'   Menggunakan reference.docx')

    result = subprocess.run(cmd, capture_output=True, text=True, cwd=SCRIPT_DIR)
    if result.returncode != 0:
        print('   pandoc stderr:', result.stderr[:500])
        raise RuntimeError('pandoc gagal! Pastikan pandoc sudah terinstall (https://pandoc.org/installing.html)')
    print(f'   ✓ DOCX dibuat: {os.path.basename(DOCX_OUT)}')


# ═════════════════════════════════════════════
# STEP 5 – Post-process DOCX
# ═════════════════════════════════════════════

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for k, v in edge_data.items():
                element.set(qn(f'w:{k}'), str(v))


def set_cell_bg(cell, color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color)
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def apply_font(run, size_pt, bold=None, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    el = rPr.find(qn('w:rFonts'))
    if el is None:
        el = OxmlElement('w:rFonts')
        rPr.insert(0, el)
    el.set(qn('w:ascii'),  FONT_NAME)
    el.set(qn('w:hAnsi'),  FONT_NAME)
    el.set(qn('w:cs'),     FONT_NAME)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


BORDER_STYLE = {"sz": "4", "val": "single", "color": BORDER_COLOR}

def step5_postprocess_docx():
    print('\n[STEP 5] Post-processing DOCX...')
    doc = Document(DOCX_OUT)

    # Default normal style
    normal = doc.styles['Normal']
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE_BODY)

    # Body paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            apply_font(run, FONT_SIZE_BODY)

    # Tables
    border_all = dict(
        top=BORDER_STYLE, bottom=BORDER_STYLE,
        start=BORDER_STYLE, end=BORDER_STYLE,
        insideH=BORDER_STYLE, insideV=BORDER_STYLE
    )
    for table in doc.tables:
        try:
            table.style = 'Table Grid'
        except Exception:
            pass

        for row_idx, row in enumerate(table.rows):
            is_header = (row_idx == 0)
            for cell in row.cells:
                set_cell_border(cell, **border_all)
                if is_header:
                    set_cell_bg(cell, HEADER_BG)
                for para in cell.paragraphs:
                    for run in para.runs:
                        apply_font(
                            run,
                            size_pt=FONT_SIZE_TABLE,
                            bold=True if is_header else None
                        )

    # Fix image sizes (max width = 14cm)
    MAX_WIDTH = Cm(14)
    RELS_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    for para in doc.paragraphs:
        for run in para.runs:
            for drawing in run._r.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
                extent = drawing.find('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                if extent is not None:
                    cx = int(extent.get('cx', 0))
                    if cx > MAX_WIDTH.emu:
                        ratio = MAX_WIDTH.emu / cx
                        cy = int(int(extent.get('cy', 0)) * ratio)
                        extent.set('cx', str(MAX_WIDTH.emu))
                        extent.set('cy', str(cy))

    doc.save(DOCX_OUT)
    print(f'   ✓ DOCX tersimpan: {os.path.basename(DOCX_OUT)}')


# ═════════════════════════════════════════════
# MAIN
# ═════════════════════════════════════════════
if __name__ == '__main__':
    print('=' * 60)
    print('  FSD New RM Sample v3.0 – DOCX Builder')
    print('=' * 60)

    step1_render_diagrams()
    step2_copy_screenshots()
    step3_preprocess_markdown()
    step4_run_pandoc()
    step5_postprocess_docx()

    # Cleanup temp file
    if os.path.exists(MD_TMP):
        os.remove(MD_TMP)
        print(f'\n   Temp file dihapus: {os.path.basename(MD_TMP)}')

    print('\n' + '=' * 60)
    print(f'  SELESAI: {DOCX_OUT}')
    print('=' * 60)
