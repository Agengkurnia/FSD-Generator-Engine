import sys
import io
import base64
import os
import re
import subprocess
import urllib.request
import zlib

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
MD_SRC      = os.path.join(SCRIPT_DIR, 'FSD_Item_Trial_Formulation_v1.0.md')
MD_TMP      = os.path.join(SCRIPT_DIR, '_tmp_item_trial_processed.md')
DOCX_OUT    = os.path.join(SCRIPT_DIR, 'FSD_Item_Trial_Formulation_v1.0.docx')
REF_DOCX    = os.path.join(SCRIPT_DIR, '..', 'NewRMSelection', 'Document', 'FSD', 'Item_Registration', 'reference.docx')
SCREENSHOTS = os.path.join(SCRIPT_DIR, 'screenshots')

os.makedirs(SCREENSHOTS, exist_ok=True)

def render_kroki(mermaid_code: str, output_path: str, label: str):
    try:
        compressed = zlib.compress(mermaid_code.strip().encode('utf-8'), 9)
        b64 = base64.urlsafe_b64encode(compressed).decode('ascii')
        url = f'https://kroki.io/mermaid/png/{b64}'
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=45) as resp:
            data = resp.read()
        with open(output_path, 'wb') as f:
            f.write(data)
        print(f'[{label}] OK -> {os.path.basename(output_path)}')
        return True
    except Exception as e:
        print(f'[{label}] Failed: {e}')
        return False

def step1_preprocess():
    print('Reading markdown...')
    with open(MD_SRC, 'r', encoding='utf-8') as f:
        text = f.read()

    blocks = list(re.finditer(r'```mermaid\s*\n(.*?)```', text, flags=re.DOTALL))
    diagram_files = []

    for i, match in enumerate(blocks):
        code = match.group(1).strip()
        png_path = os.path.join(SCREENSHOTS, f'diagram_{i}.png')
        render_kroki(code, png_path, f'Diagram_{i}')
        diagram_files.append(png_path)

    diagram_counter = [0]
    def replace_mermaid(m):
        path = diagram_files[diagram_counter[0]]
        diagram_counter[0] += 1
        rel_path = os.path.relpath(path, SCRIPT_DIR).replace('\\', '/')
        if os.path.exists(path):
            return f'\n![Diagram]({rel_path})\n'
        return '\n>[Diagram render failed]\n'

    text = re.sub(r'```mermaid\s*\n(.*?)```', replace_mermaid, text, flags=re.DOTALL)

    with open(MD_TMP, 'w', encoding='utf-8') as f:
        f.write(text)

def step2_pandoc():
    print('Running pandoc...')
    cmd = [
        'pandoc', MD_TMP,
        '-o', DOCX_OUT,
        '--from=markdown+pipe_tables',
        f'--resource-path={SCRIPT_DIR};{SCREENSHOTS}',
        '--toc',
        '--toc-depth=3',
    ]
    if os.path.exists(REF_DOCX):
        cmd += [f'--reference-doc={REF_DOCX}']

    res = subprocess.run(cmd, capture_output=True, text=True, cwd=SCRIPT_DIR)
    if res.returncode != 0:
        print(res.stderr)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = f'w:{edge}'
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for k, v in kwargs[edge].items():
                element.set(qn(f'w:{k}'), str(v))

def set_cell_bg(cell, color: str):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)

def apply_font(run, size_pt, bold=None):
    run.font.name = 'Calibri'
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    el = rPr.find(qn('w:rFonts'))
    if el is None:
        el = OxmlElement('w:rFonts')
        rPr.insert(0, el)
    el.set(qn('w:ascii'), 'Calibri')
    el.set(qn('w:hAnsi'), 'Calibri')
    el.set(qn('w:cs'), 'Calibri')
    if bold is not None: run.bold = bold

def step3_postprocess():
    print('Postprocessing DOCX...')
    if not os.path.exists(DOCX_OUT): return
    doc = Document(DOCX_OUT)

    bdr = {"sz": "8", "val": "single", "color": "000000"}
    b_all = { 'top':bdr,'bottom':bdr,'start':bdr,'end':bdr,'insideH':bdr,'insideV':bdr }

    for t in doc.tables:
        try: t.style = 'Table Grid'
        except: pass
        for ri, row in enumerate(t.rows):
            is_hdr = (ri == 0)
            for cell in row.cells:
                set_cell_border(cell, **b_all)
                if is_hdr: set_cell_bg(cell, 'D9EAD3')
                for p in cell.paragraphs:
                    for r in p.runs:
                        apply_font(r, 9, bold=True if is_hdr else None)

    MAX_W = Cm(15)
    for p in doc.paragraphs:
        for r in p.runs:
            apply_font(r, 11)
            dw_ins = r._r.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
            for dw in dw_ins:
                ext = dw.find('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                if ext is not None:
                    cx = int(ext.get('cx', 0))
                    if cx > MAX_W.emu:
                        ratio = MAX_W.emu / cx
                        cy = int(int(ext.get('cy', 0)) * ratio)
                        ext.set('cx', str(MAX_W.emu))
                        ext.set('cy', str(cy))

    doc.save(DOCX_OUT)

if __name__ == '__main__':
    step1_preprocess()
    step2_pandoc()
    step3_postprocess()
    if os.path.exists(MD_TMP): os.remove(MD_TMP)
    print('Done!')
