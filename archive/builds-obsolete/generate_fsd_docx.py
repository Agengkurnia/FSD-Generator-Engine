"""
FSD DOCX Generator v2 - with Mermaid diagram rendering
- Renders Mermaid diagrams to PNG via mermaid.ink API
- Replaces mermaid code blocks in markdown with image references
- Generates DOCX from patched markdown using pandoc
- Applies table formatting: borders + yellow header background
"""

import subprocess
import sys
import os
import re
import base64
import urllib.request
import shutil
from pathlib import Path

# Paths
FSD_MD = r"C:\Users\Lenovo\.gemini\antigravity\brain\2e88ab95-8a65-4966-82f2-b215afb1f06f\FSD_New_RM_Sample.md"
OUTPUT_DOCX = r"d:\Work\Source\RMSelection Hijrah\NewRMSelection\Document\FSD\FSD_New_RM_Sample_v1.8.docx"
SCREENSHOTS_DIR = r"C:\Users\Lenovo\.gemini\antigravity\brain\2e88ab95-8a65-4966-82f2-b215afb1f06f\screenshots"
TEMP_MD = r"d:\Work\Source\RMSelection Hijrah\NewRMSelection\Document\FSD\_temp_fsd_rendered.md"

os.makedirs(SCREENSHOTS_DIR, exist_ok=True)


def render_mermaid_to_png(mermaid_code: str, output_path: str, diagram_name: str) -> bool:
    """Render Mermaid code to PNG using mermaid.ink API, with cache check"""
    # Use cached PNG if it exists
    if os.path.exists(output_path):
        size_kb = os.path.getsize(output_path) / 1024
        print(f"  [CACHE] '{diagram_name}' already exists ({size_kb:.0f} KB), skipping render")
        return True

    try:
        # Encode the mermaid code to base64
        encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
        url = f"https://mermaid.ink/img/{encoded}?bgColor=white&width=900"

        print(f"  Rendering '{diagram_name}' via mermaid.ink ...")
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=30) as response:
            png_data = response.read()

        with open(output_path, 'wb') as f:
            f.write(png_data)

        size_kb = len(png_data) / 1024
        print(f"  [OK] Saved: {output_path} ({size_kb:.0f} KB)")
        return True

    except Exception as e:
        print(f"  [WARN] Failed to render '{diagram_name}': {e}")
        return False


def step0_render_mermaid_diagrams(content: str) -> str:
    """
    Find all ```mermaid blocks, render to PNG, replace with image markdown.
    Returns patched markdown content.
    """
    print("Step 0: Rendering Mermaid diagrams to PNG images ...")

    mermaid_pattern = re.compile(r'```mermaid\n(.*?)```', re.DOTALL)
    matches = list(mermaid_pattern.finditer(content))

    print(f"  Found {len(matches)} Mermaid diagram(s)")

    rendered_count = 0
    offset = 0  # track string position shift as we replace

    patched = content
    replacements = []

    for i, match in enumerate(matches):
        mermaid_code = match.group(1).strip()
        diagram_name = f"FSD_Diagram_{i+1:02d}"
        png_path = os.path.join(SCREENSHOTS_DIR, f"{diagram_name}.png")

        success = render_mermaid_to_png(mermaid_code, png_path, diagram_name)
        if success:
            # Build replacement: image markdown (uses forward slashes for pandoc compat)
            png_fwd = png_path.replace('\\', '/')
            img_md = f"![]({png_fwd})"
            replacements.append((match.start(), match.end(), img_md))
            rendered_count += 1
        else:
            print(f"  [INFO] Keeping code block for diagram {i+1} (render failed)")

    # Apply replacements in reverse order to preserve positions
    for start, end, img_md in reversed(replacements):
        patched = patched[:start] + img_md + patched[end:]

    print(f"  [OK] Rendered {rendered_count}/{len(matches)} diagrams (rest kept as code)")
    return patched


def step1_generate_docx(md_content_patched: str) -> bool:
    """Write patched markdown to temp file, then generate DOCX from it using pandoc"""
    print("\nStep 1: Generating DOCX from Markdown with pandoc ...")

    with open(TEMP_MD, 'w', encoding='utf-8') as f:
        f.write(md_content_patched)

    cmd = [
        "pandoc", TEMP_MD,
        "-o", OUTPUT_DOCX,
        "--from=markdown",
        "--to=docx",
        "--toc",
        "--toc-depth=3",
        f"--resource-path={SCREENSHOTS_DIR}"
    ]

    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"  [ERROR] pandoc failed:\n{result.stderr}")
        return False

    if result.stderr:
        warnings = [l for l in result.stderr.split('\n') if 'WARNING' in l and 'does not exist' in l]
        if warnings:
            print(f"  [!] {len(warnings)} image(s) not found (will show as broken)")

    print(f"  [OK] DOCX generated: {OUTPUT_DOCX}")

    # Clean up temp file
    try:
        os.remove(TEMP_MD)
    except Exception:
        pass

    return True


def step2_format_tables() -> bool:
    """Apply table formatting: borders + yellow headers"""
    print("\nStep 2: Applying table formatting ...")

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        print("  ERROR: python-docx not installed")
        print("  Run: python -m pip install python-docx")
        return False

    doc = Document(OUTPUT_DOCX)

    table_count = 0
    for table in doc.tables:
        table_count += 1
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

        borders_xml = f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        </w:tblBorders>
        '''

        existing_borders = tbl_pr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tbl_pr.remove(existing_borders)
        tbl_pr.append(parse_xml(borders_xml))

        if len(table.rows) > 0:
            header_row = table.rows[0]
            for cell in header_row.cells:
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="FFD700" w:val="clear"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)

        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:
                continue
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    doc.save(OUTPUT_DOCX)
    print(f"  [OK] Formatted {table_count} tables (borders + yellow headers)")
    return True


def main():
    print("=" * 60)
    print("FSD DOCX Generator (with Mermaid rendering)")
    print("=" * 60)
    print()

    # Read markdown
    with open(FSD_MD, 'r', encoding='utf-8') as f:
        content = f.read()

    # Step 0: Render Mermaid to PNG and patch markdown
    patched_content = step0_render_mermaid_diagrams(content)

    # Step 1: Generate DOCX from patched markdown
    if not step1_generate_docx(patched_content):
        sys.exit(1)

    # Step 2: Format tables
    if not step2_format_tables():
        sys.exit(1)

    print()
    print("=" * 60)
    print(f"[OK] DONE! File: {OUTPUT_DOCX}")
    file_size = os.path.getsize(OUTPUT_DOCX)
    print(f"  Size: {file_size / 1024:.0f} KB")
    print("=" * 60)


if __name__ == "__main__":
    main()
