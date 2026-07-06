"""
build_new_rm_sample_docx.py
Full pipeline to build a styled FSD DOCX for New RM Sample Management:
  1. Render mermaid diagram → PNG via mermaid.ink
  2. Pre-process markdown: replace ```mermaid blocks and *(Gambar N: ...)* placeholders with ![...](path)
  3. Run pandoc (with reference.docx for Calibri styles)
  4. Post-process DOCX: add table borders + yellow header rows

Run: python build_new_rm_sample_docx.py
"""

import base64
import os
import re
import subprocess
import urllib.request

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────────
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
MD_IN       = os.path.join(SCRIPT_DIR, 'FSD_New_RM_Sample_v1.9.md')
MD_TMP      = os.path.join(SCRIPT_DIR, '_tmp_new_rm_sample_processed.md')
DOCX_OUT    = os.path.join(SCRIPT_DIR, 'FSD_New_RM_Sample_v1.9_Update.docx')
REF_DOCX    = os.path.join(SCRIPT_DIR, 'reference.docx')
SCREENSHOTS = os.path.join(SCRIPT_DIR, 'screenshots')

# ── Mermaid Flow Diagram ───────────────────────────────────────────────────────
FLOW_MERMAID_CODE = """
flowchart TD
    subgraph Row1 [Tahap 1: Evaluasi]
        direction LR
        Start(["Mulai"]) --> Step1["STEP 1: DOCUMENT SAMPLE"] --> Step2["STEP 2: SAMPLE PURPOSE"] --> Step3["STEP 3: RM EVALUATION"]
    end
    
    Step3 --> Step4["STEP 4: DISPOSITION"]
    
    subgraph Row2 [Tahap 2: Keputusan]
        direction RL
        Step4 --> Decision{"Keputusan?"}
        Decision --> |"APPROVED"| ApprovalCheck{"Perlu Approval?"}
        Decision --> |"REJECTED"| EndReject(["REJECTED"])
        Decision --> |"ON HOLD"| OnHold["ON HOLD"]
        Decision --> |"CANCELLED"| EndCancel(["CANCELLED"])
    end
    
    OnHold -.-> |"Kembali"| Step3
    ApprovalCheck --> PendingApproval["Pending Approval"]
    ApprovalCheck --> |"Tidak"| Completed1(["COMPLETED"])
    
    subgraph Row3 [Tahap 3: Approval Atasan]
        direction LR
        PendingApproval --> ApproverReview["Approver Review"] --> ApprovalDecision{"Hasil Approval?"}
        ApprovalDecision --> |"APPROVE"| Completed(["COMPLETED"])
    end
    
    ApprovalDecision -.-> |"REJECT"| Step4

    style Start fill:#8CC63F,color:#fff,stroke:#009B4D,stroke-width:2px
    style Completed fill:#388E3C,color:#fff,stroke:#1B5E20,stroke-width:2px
    style Completed1 fill:#388E3C,color:#fff,stroke:#1B5E20,stroke-width:2px
    style EndReject fill:#D32F2F,color:#fff,stroke:#B71C1C,stroke-width:2px
    style EndCancel fill:#9E9E9E,color:#fff,stroke:#616161,stroke-width:2px
    style Step1 fill:#FFF3CD,stroke:#FFC107,stroke-width:2px
    style Step2 fill:#D1ECF1,stroke:#17A2B8,stroke-width:2px
    style Step3 fill:#E2E3E5,stroke:#6C757D,stroke-width:2px
    style Step4 fill:#F8D7DA,stroke:#DC3545,stroke-width:2px
    style Decision fill:#FFF3E0,stroke:#F7941E,stroke-width:3px
    style ApprovalCheck fill:#FFF3E0,stroke:#F7941E,stroke-width:2px
    style ApprovalDecision fill:#FFF3E0,stroke:#F7941E,stroke-width:2px
    style OnHold fill:#E7F3FF,stroke:#0D6EFD,stroke-width:2px
    style PendingApproval fill:#FFF9C4,stroke:#F57F17,stroke-width:2px
    style ApproverReview fill:#F3E5F5,stroke:#7B1FA2,stroke-width:2px
"""
FLOW_PNG = os.path.join(SCREENSHOTS, 'new_rm_sample_flow.png')

# ── ERD Diagram ────────────────────────────────────────────────────────────────
ERD_MERMAID_CODE = """
erDiagram
    T_RM_SAMPLE_HEADER {
        BIGSERIAL  intSampleId       PK
        VARCHAR50  txtSampleNo       UK
        BIGINT     intTypeDocumentId FK
        VARCHAR200 txtMaterialName
        BIGINT     intSupplierId     FK
        VARCHAR200 txtSupplierName
        DECIMAL184 decPrice
        BIGINT     intCurrencyId     FK
        BIGINT     intUOMId          FK
        DECIMAL184 decQuantity
        BIGINT     intPICId          FK
        TIMESTAMP  dtSubmissionDate
        VARCHAR50  txtWorkflowStage
        VARCHAR50  txtStatus
        BOOLEAN    bitActive
    }
    T_RM_SAMPLE_ATTACHMENT {
        BIGSERIAL  intAttachId    PK
        BIGINT     intSampleId   FK
        BIGINT     intDocTypeId  FK
        VARCHAR255 txtFileName
        TEXT       txtFilePath
        VARCHAR50  txtUploadedBy
        TIMESTAMP  dtUploadedDate
    }
    T_RM_SAMPLE_PURPOSE {
        BIGSERIAL  intPurposeId         PK
        BIGINT     intSampleId          FK
        BIGINT     intPurposeCategoryId FK
        TEXT       txtTargetProduct
        VARCHAR1000 txtExpectedBenefit
        DECIMAL52  decUsagePercentage
        DATE       dtAnalysisDeadline
    }
    T_RM_SAMPLE_EVALUATION {
        BIGSERIAL  intEvalId          PK
        BIGINT     intSampleId        FK
        DATE       dtTestDate
        BIGINT     intTestTypeId      FK
        VARCHAR10  txtTestResult
        VARCHAR10  txtSpecCompliance
        INT        intQualityScore
        TEXT       txtEvaluationNotes
        BIGINT     intEvaluatorId     FK
    }
    T_RM_SAMPLE_DISPOSITION {
        BIGSERIAL  intDispId             PK
        BIGINT     intSampleId           FK
        VARCHAR20  txtDecision
        TIMESTAMP  dtDecisionDate
        BIGINT     intDecisionById       FK
        VARCHAR1000 txtReason
        BOOLEAN    bitApprovalRequired
        BIGINT     intApproverId         FK
        VARCHAR20  txtApprovalStatus
    }

    T_RM_SAMPLE_HEADER     ||--o{ T_RM_SAMPLE_ATTACHMENT   : "has Attachments"
    T_RM_SAMPLE_HEADER     ||--o| T_RM_SAMPLE_PURPOSE      : "has Purpose"
    T_RM_SAMPLE_HEADER     ||--o{ T_RM_SAMPLE_EVALUATION   : "has Evaluations"
    T_RM_SAMPLE_HEADER     ||--o| T_RM_SAMPLE_DISPOSITION  : "has Disposition"
"""

ERD_PNG = os.path.join(SCREENSHOTS, 'new_rm_sample_erd.png')

# ── Screenshot placeholders ────────────────────────────────────────────────────
SCREENSHOT_MAP = {
    'Gambar 1': os.path.join(SCREENSHOTS, 'NewRMSample_Index_Page.png'),
    'Gambar 2': os.path.join(SCREENSHOTS, 'NewRMSample_Detail_Wizard.png'),
    'Gambar 3': os.path.join(SCREENSHOTS, 'Step1_Document_Sample.png'),
    'Gambar 4': os.path.join(SCREENSHOTS, 'Step2_Sample_Purpose.png'),
    'Gambar 5': os.path.join(SCREENSHOTS, 'Step3_RM_Evaluation.png'),
    'Gambar 6': os.path.join(SCREENSHOTS, 'Step4_Disposition.png'),
    'Gambar 7': ERD_PNG,
}

# ── Styling constants ──────────────────────────────────────────────────────────
HEADER_BG    = '9BC2E6'   # light blue
BORDER_COLOR = '000000'   # black
FONT_NAME    = 'Calibri'
FONT_SIZE    = 11


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 – Render mermaid diagrams to PNG
# ══════════════════════════════════════════════════════════════════════════════
def render_mermaid_png(mermaid_code, output_path, label):
    os.makedirs(SCREENSHOTS, exist_ok=True)
    encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('ascii')
    url = f'https://mermaid.ink/img/{encoded}?type=png&scale=3'
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = resp.read()
        with open(output_path, 'wb') as f:
            f.write(data)
        print(f'   [{label}] Saved {len(data):,} bytes -> {output_path}')
    except Exception as e:
        print(f'   [{label}] Warning: mermaid.ink failed ({e}). Using existing PNG if available.')


def render_mermaid():
    print('[1/4] Rendering mermaid diagrams...')
    # Use locally generated high-res flowchart instead of mermaid.ink
    # render_mermaid_png(FLOW_MERMAID_CODE, FLOW_PNG, 'Flow')
    render_mermaid_png(ERD_MERMAID_CODE, ERD_PNG, 'ERD')


# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 – Pre-process markdown
# ══════════════════════════════════════════════════════════════════════════════
def preprocess_markdown():
    print('[2/4] Pre-processing markdown...')
    with open(MD_IN, 'r', encoding='utf-8') as f:
        text = f.read()

    # Replace ```mermaid ... ``` blocks with inline image reference
    def replace_mermaid(m):
        if os.path.exists(FLOW_PNG):
            rel = os.path.relpath(FLOW_PNG, SCRIPT_DIR).replace('\\', '/')
            return f'\n![New RM Sample Business Flow Diagram]({rel})\n'
        return '*(Flow Diagram tidak tersedia)*\n'

    text = re.sub(r'```mermaid.*?```', replace_mermaid, text, flags=re.DOTALL)

    # Replace *(Gambar N: ...)*  with inline image reference
    def replace_screenshot(m):
        label   = m.group(1).strip()   # e.g. "Gambar 1"
        caption = m.group(2).strip()   # e.g. "Halaman Index"
        for key, path in SCREENSHOT_MAP.items():
            if key in label and os.path.exists(path):
                rel = os.path.relpath(path, SCRIPT_DIR).replace('\\', '/')
                return f'\n![{caption}]({rel})\n'
        # If screenshot not found, keep a visible note
        return f'\n*({label}: {caption} — screenshot belum tersedia)*\n'

    text = re.sub(r'\*\(?(Gambar\s*\d+)[:\s]*([^)]*)\)?\*', replace_screenshot, text)

    with open(MD_TMP, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f'   Written: {MD_TMP}')


# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 – Run pandoc
# ══════════════════════════════════════════════════════════════════════════════
def run_pandoc():
    print('[3/4] Running pandoc...')
    cmd = [
        'pandoc', MD_TMP,
        '-o', DOCX_OUT,
        '--from=markdown+pipe_tables',
        f'--resource-path={SCRIPT_DIR}',
        '--toc',
        '--toc-depth=3',
    ]
    if os.path.exists(REF_DOCX):
        cmd += [f'--reference-doc={REF_DOCX}']
        print(f'   Using reference doc: {REF_DOCX}')
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print('   pandoc stderr:', result.stderr)
        raise RuntimeError('pandoc failed')
    print(f'   Generated: {DOCX_OUT}')


# ══════════════════════════════════════════════════════════════════════════════
# STEP 4 – Post-process DOCX (tables, fonts)
# ══════════════════════════════════════════════════════════════════════════════
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_cell_bg_color(cell, color_str):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_str)
    cell._tc.get_or_add_tcPr().append(shading)


def apply_font(run, size_pt, bold=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    el = rPr.find(qn('w:rFonts'))
    if el is None:
        el = OxmlElement('w:rFonts')
        rPr.insert(0, el)
    el.set(qn('w:ascii'),  FONT_NAME)
    el.set(qn('w:hAnsi'), FONT_NAME)
    el.set(qn('w:cs'),    FONT_NAME)
    if bold is not None:
        run.bold = bold


def postprocess_docx():
    print('[4/4] Post-processing DOCX (borders, colours, fonts)...')
    doc = Document(DOCX_OUT)

    # Set default text style to Calibri
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE)

    # Body paragraphs – apply Calibri
    for para in doc.paragraphs:
        for run in para.runs:
            apply_font(run, FONT_SIZE)

    # Tables
    for table in doc.tables:
        try:
            table.style = 'Table Grid'
        except Exception:
            pass
        for row_idx, row in enumerate(table.rows):
            is_header = (row_idx == 0)
            for cell in row.cells:
                set_cell_border(
                    cell,
                    top={"sz": 4, "val": "single", "color": "000000"},
                    bottom={"sz": 4, "val": "single", "color": "000000"},
                    start={"sz": 4, "val": "single", "color": "000000"},
                    end={"sz": 4, "val": "single", "color": "000000"}
                )
                if is_header:
                    set_cell_bg_color(cell, HEADER_BG)
                
                for para in cell.paragraphs:
                    for run in para.runs:
                        apply_font(run, size_pt=10, bold=True if is_header else None)

    doc.save(DOCX_OUT)
    print(f'   Saved styled DOCX -> {DOCX_OUT}')


# ══════════════════════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print('=' * 60)
    print('  Building FSD New RM Sample v1.8 DOCX')
    print('=' * 60)
    render_mermaid()
    preprocess_markdown()
    run_pandoc()
    postprocess_docx()
    # Cleanup temp markdown
    if os.path.exists(MD_TMP):
        os.remove(MD_TMP)
        print(f'   Cleaned up temp: {MD_TMP}')
    print()
    print('Done! Output file:', DOCX_OUT)
