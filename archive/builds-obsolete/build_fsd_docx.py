"""
build_fsd_docx.py
Full pipeline to build a styled FSD DOCX:
  1. Render mermaid diagram → PNG via mermaid.ink
  2. Pre-process markdown: replace ```mermaid blocks and *(Gambar N: ...)* placeholders with ![...](path)
  3. Run pandoc (with reference.docx for Calibri styles)
  4. Post-process DOCX: add table borders + light-green header rows

Run: python build_fsd_docx.py
"""

import base64
import os
import re
import shutil
import subprocess
import urllib.request

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────────
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
MD_IN       = os.path.join(SCRIPT_DIR, 'FSD_SEAL_Project_Identity_v1.0.md')
MD_TMP      = os.path.join(SCRIPT_DIR, '_tmp_fsd_processed.md')
DOCX_OUT    = os.path.join(SCRIPT_DIR, 'FSD_SEAL_Project_Identity_v1.0.docx')
REF_DOCX    = os.path.join(SCRIPT_DIR, 'reference.docx')
SCREENSHOTS = os.path.join(SCRIPT_DIR, 'screenshots')

# ── Mermaid diagrams ───────────────────────────────────────────────────────────
MERMAID_CODE = """
flowchart TD
    A(["Buat SEAL Baru"]) --> B["Input Data SEAL Identity (Product Group, Supplier, Item Code, PIC)"]
    B --> C["Document Upload: PIC upload file per jenis dokumen sesuai PM Category"]
    C --> D{{Submit?}}
    D -->|Ya| E["Status: WAIT_APPROVAL"]
    E --> F{{"Semua Reviewer Approve? (AND Logic)"}}
    F -->|"Semua Approve"| G["Status: APPROVED"]
    G --> K(["Status: COMPLETED"])
    F -->|"Salah satu Revise"| I["Status: REVISE"]
    I -->|"Notifikasi ke Uploader - Upload Ulang"| C

    style A fill:#4CAF50,color:#fff
    style K fill:#1565C0,color:#fff
    style G fill:#4CAF50,color:#fff
    style I fill:#F44336,color:#fff
    style E fill:#FF9800,color:#fff
    style F fill:#FF9800,color:#fff
"""

DIAGRAM_PNG = os.path.join(SCREENSHOTS, 'seal_flow_diagram.png')

# ── ERD diagram ────────────────────────────────────────────────────────────────
ERD_CODE = """
erDiagram
    trIdcSealHdr {
        BIGSERIAL   IntId       PK
        VARCHAR50   TxtSealNo   UK
        VARCHAR100  TxtProductGroup
        VARCHAR100  TxtProjectType
        VARCHAR10   TxtPotsType
        VARCHAR255  TxtSupplierName
        VARCHAR100  TxtItemCodePm
        VARCHAR500  TxtItemDesc
        VARCHAR100  TxtAkasiaNumber
        VARCHAR100  TxtMdNumber
        DATE        DtMdExpiry
        VARCHAR100  TxtItemSpec
        VARCHAR100  TxtPmEvaluation
        VARCHAR30   TxtDocStatus
        VARCHAR100  TxtCreatedBy
        TIMESTAMP   DtCreated
        VARCHAR100  TxtUpdatedBy
        TIMESTAMP   DtUpdated
    }
    trIdcSealPic {
        BIGSERIAL   IntId       PK
        BIGINT      IntHdrId    FK
        VARCHAR100  TxtUsername
        VARCHAR100  TxtKnGlobal
        VARCHAR100  TxtDepartment
        BOOLEAN     BitActive
        VARCHAR100  TxtCreatedBy
        TIMESTAMP   DtCreated
    }
    trIdcSealDoc {
        BIGSERIAL   IntId       PK
        BIGINT      IntHdrId    FK
        VARCHAR100  TxtDocType
        VARCHAR255  TxtDocName
        TEXT        TxtFilePath
        VARCHAR20   TxtFileExt
        VARCHAR30   TxtDocStatus
        VARCHAR100  TxtUploadedBy
        TIMESTAMP   DtUploaded
        VARCHAR100  TxtCreatedBy
        TIMESTAMP   DtCreated
    }
    trIdcSealReviewer {
        BIGSERIAL   IntId       PK
        BIGINT      IntHdrId    FK
        VARCHAR100  TxtReviewerName
        VARCHAR100  TxtReviewerRole
        VARCHAR30   TxtReviewStatus
        TEXT        TxtReviewNote
        TIMESTAMP   DtReviewed
        VARCHAR100  TxtCreatedBy
        TIMESTAMP   DtCreated
    }
    trIdcSealStatusLog {
        BIGSERIAL   IntId       PK
        BIGINT      IntHdrId    FK
        VARCHAR30   TxtStatusFrom
        VARCHAR30   TxtStatusTo
        TEXT        TxtNote
        VARCHAR100  TxtActionBy
        TIMESTAMP   DtAction
    }
    trPmCategoryDoc {
        BIGSERIAL   IntId       PK
        VARCHAR100  TxtPmCategory
        VARCHAR255  TxtDocName
        BOOLEAN     BitActive
        VARCHAR100  TxtCreatedBy
        TIMESTAMP   DtCreated
    }
    mPMCategoryDocHeader {
        INT         IntPMCategoryDocHdrId   PK
        INT         IntPMCategoryId         FK
        VARCHAR     TxtPMCategoryCode
        VARCHAR     TxtPMCategoryName
        BOOLEAN     BitActive
    }
    mPMCategoryDocDetail {
        INT         IntPMCategoryDocDtlId    PK
        INT         IntPMCategoryDocHdrId    FK
        VARCHAR255  TxtDocumentName
        BOOLEAN     BitActive
    }

    trIdcSealHdr     ||--o{ trIdcSealPic        : "has PIC"
    trIdcSealHdr     ||--o{ trIdcSealDoc        : "has Documents"
    trIdcSealHdr     ||--o{ trIdcSealReviewer   : "has Reviewers"
    trIdcSealHdr     ||--o{ trIdcSealStatusLog  : "has Status Log"
    mPMCategoryDocHeader ||--o{ mPMCategoryDocDetail : "defines Docs"
    mPMCategoryDocDetail ||--o{ trIdcSealDoc    : "template for"
"""

ERD_PNG = os.path.join(SCREENSHOTS, 'seal_erd.png')

# ── Screenshot placeholders ────────────────────────────────────────────────────
SCREENSHOT_MAP = {
    'Gambar 1': os.path.join(SCREENSHOTS, 'ss_index.png'),
    'Gambar 2': os.path.join(SCREENSHOTS, 'ss_detail.png'),
    'Gambar 3': os.path.join(SCREENSHOTS, 'ss_approval.png'),
    'Gambar 4': os.path.join(SCREENSHOTS, 'ss_approval_detail.png'),
    'Gambar 5': ERD_PNG,
}

# ── Styling constants ──────────────────────────────────────────────────────────
HEADER_BG    = 'C6EFCE'   # light green
BORDER_COLOR = '538135'   # dark green
FONT_NAME    = 'Calibri'
FONT_SIZE    = 11


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 – Render mermaid diagrams to PNG
# ══════════════════════════════════════════════════════════════════════════════
def render_mermaid_png(mermaid_code, output_path, label):
    os.makedirs(SCREENSHOTS, exist_ok=True)
    encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('ascii')
    url = f'https://mermaid.ink/img/{encoded}?type=png'
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = resp.read()
        with open(output_path, 'wb') as f:
            f.write(data)
        print(f'   [{label}] Saved {len(data):,} bytes -> {output_path}')
    except Exception as e:
        print(f'   [{label}] Warning: mermaid.ink failed ({e}). Using existing PNG if available.')


def render_mermaid():
    print('[1/4] Rendering mermaid diagrams...')
    render_mermaid_png(MERMAID_CODE, DIAGRAM_PNG, 'Flow')
    render_mermaid_png(ERD_CODE, ERD_PNG, 'ERD')


# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 – Pre-process markdown
# ══════════════════════════════════════════════════════════════════════════════
def preprocess_markdown():
    print('[2/4] Pre-processing markdown...')
    with open(MD_IN, 'r', encoding='utf-8') as f:
        text = f.read()

    # Replace ```mermaid ... ``` blocks with inline image reference
    def replace_mermaid(m):
        if os.path.exists(DIAGRAM_PNG):
            rel = os.path.relpath(DIAGRAM_PNG, SCRIPT_DIR).replace('\\', '/')
            return f'\n![SEAL Business Flow Diagram]({rel})\n'
        return '*(Diagram tidak tersedia)*\n'

    text = re.sub(r'```mermaid.*?```', replace_mermaid, text, flags=re.DOTALL)

    # Replace *(Gambar N: ...)*  with inline image reference
    def replace_screenshot(m):
        label = m.group(1).strip()   # e.g. "Gambar 1"
        caption = m.group(2).strip() # e.g. "Halaman Index SEAL"
        for key, path in SCREENSHOT_MAP.items():
            if key in label and os.path.exists(path):
                rel = os.path.relpath(path, SCRIPT_DIR).replace('\\', '/')
                return f'\n![{caption}]({rel})\n'
        return m.group(0)  # keep original if not found

    text = re.sub(r'\*\(?(Gambar\s*\d+)[:\s]*([^)]*)\)?\*', replace_screenshot, text)

    with open(MD_TMP, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f'   Written: {MD_TMP}')


# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 – Run pandoc
# ══════════════════════════════════════════════════════════════════════════════
def run_pandoc():
    print('[3/4] Running pandoc...')
    cmd = [
        'pandoc', MD_TMP,
        '-o', DOCX_OUT,
        '--from=markdown+pipe_tables',
        f'--resource-path={SCRIPT_DIR}',
        '--toc',
        '--toc-depth=3',
    ]
    if os.path.exists(REF_DOCX):
        cmd += [f'--reference-doc={REF_DOCX}']
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print('   pandoc stderr:', result.stderr)
        raise RuntimeError('pandoc failed')
    print(f'   Generated: {DOCX_OUT}')


# ══════════════════════════════════════════════════════════════════════════════
# STEP 4 – Post-process DOCX (tables, fonts)
# ══════════════════════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color, sz=8):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def apply_font(run, bold=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    rPr = run._r.get_or_add_rPr()
    el = rPr.find(qn('w:rFonts'))
    if el is None:
        el = OxmlElement('w:rFonts')
        rPr.insert(0, el)
    el.set(qn('w:ascii'), FONT_NAME)
    el.set(qn('w:hAnsi'), FONT_NAME)
    el.set(qn('w:cs'), FONT_NAME)
    if bold is not None:
        run.bold = bold


def postprocess_docx():
    print('[4/4] Post-processing DOCX (borders, colours, fonts)...')
    doc = Document(DOCX_OUT)

    # Body paragraphs – apply Calibri
    for para in doc.paragraphs:
        for run in para.runs:
            apply_font(run)

    # Tables
    for table in doc.tables:
        try:
            table.style = doc.styles['Table Grid']
        except Exception:
            pass
        for ri, row in enumerate(table.rows):
            is_header = (ri == 0)
            for cell in row.cells:
                set_cell_borders(cell, BORDER_COLOR)
                if is_header:
                    set_cell_bg(cell, HEADER_BG)
                for para in cell.paragraphs:
                    for run in para.runs:
                        apply_font(run, bold=True if is_header else None)

    doc.save(DOCX_OUT)
    print(f'   Saved styled DOCX -> {DOCX_OUT}')


# ══════════════════════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    render_mermaid()
    preprocess_markdown()
    run_pandoc()
    postprocess_docx()
    # Cleanup temp markdown
    if os.path.exists(MD_TMP):
        os.remove(MD_TMP)
    print('\nDone! Open:', DOCX_OUT)
