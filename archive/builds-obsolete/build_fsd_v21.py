"""
build_fsd_v2.py
Build pipeline for FSD New RM Sample v2.0 DOCX:
  1. Render mermaid diagrams → PNG via mermaid.ink
  2. Pre-process markdown: replace ```mermaid blocks and *(Gambar N: ...)* with ![...](path)
  3. Run pandoc (with reference.docx)
  4. Post-process DOCX: borders + blue header rows

Run: python build_fsd_v2.py
"""

import base64
import os
import re
import subprocess
import urllib.request

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────────
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
MD_IN       = os.path.join(SCRIPT_DIR, 'FSD_New_RM_Sample_v2.1.md')
MD_TMP      = os.path.join(SCRIPT_DIR, '_tmp_v2_processed.md')
DOCX_OUT    = os.path.join(SCRIPT_DIR, 'FSD_New_RM_Sample_v2.1.docx')
REF_DOCX    = os.path.join(SCRIPT_DIR, 'reference.docx')
SCREENSHOTS = os.path.join(SCRIPT_DIR, 'screenshots')

# ── Mermaid Flow Diagram ───────────────────────────────────────────────────────
FLOW_MERMAID_CODE = """
flowchart TD
    subgraph Row1 [Tahap 1: Evaluasi]
        direction LR
        Start(["Mulai"]) --> Step1["STEP 1: DOCUMENT REGISTRATION"] --> Step2["STEP 2: DOCUMENT SAMPLE"] --> Step3["STEP 3: RM EVALUATION"]
    end
    
    Step3 --> Step4["STEP 4: DISPOSITION"]
    
    subgraph Row2 [Tahap 2: Keputusan]
        direction RL
        Step4 --> Decision{"Keputusan?"}
        Decision --> |"APPROVED"| ApprovalCheck{"Perlu Approval?"}
        Decision --> |"REJECTED"| EndReject(["REJECTED"])
        Decision --> |"ON HOLD"| OnHold["ON HOLD"]
        Decision --> |"CANCELLED"| EndCancel(["CANCELLED"])
    end
    
    OnHold -.-> |"Kembali"| Step3
    ApprovalCheck --> PendingApproval["Pending Approval"]
    ApprovalCheck --> |"Tidak"| Completed1(["COMPLETED"])
    
    subgraph Row3 [Tahap 3: Approval Atasan]
        direction LR
        PendingApproval --> ApproverReview["Approver Review"] --> ApprovalDecision{"Hasil Approval?"}
        ApprovalDecision --> |"APPROVE"| Completed(["COMPLETED"])
    end
    
    ApprovalDecision -.-> |"REJECT"| Step4

    style Start fill:#8CC63F,color:#fff,stroke:#009B4D,stroke-width:2px
    style Completed fill:#388E3C,color:#fff,stroke:#1B5E20,stroke-width:2px
    style Completed1 fill:#388E3C,color:#fff,stroke:#1B5E20,stroke-width:2px
    style EndReject fill:#D32F2F,color:#fff,stroke:#B71C1C,stroke-width:2px
    style EndCancel fill:#9E9E9E,color:#fff,stroke:#616161,stroke-width:2px
    style Step1 fill:#FFF3CD,stroke:#FFC107,stroke-width:2px
    style Step2 fill:#D1ECF1,stroke:#17A2B8,stroke-width:2px
    style Step3 fill:#E2E3E5,stroke:#6C757D,stroke-width:2px
    style Step4 fill:#F8D7DA,stroke:#DC3545,stroke-width:2px
    style Decision fill:#FFF3E0,stroke:#F7941E,stroke-width:3px
    style ApprovalCheck fill:#FFF3E0,stroke:#F7941E,stroke-width:2px
    style ApprovalDecision fill:#FFF3E0,stroke:#F7941E,stroke-width:2px
    style OnHold fill:#E7F3FF,stroke:#0D6EFD,stroke-width:2px
    style PendingApproval fill:#FFF9C4,stroke:#F57F17,stroke-width:2px
    style ApproverReview fill:#F3E5F5,stroke:#7B1FA2,stroke-width:2px
"""
FLOW_PNG = os.path.join(SCREENSHOTS, 'new_rm_sample_flow_v2.png')

# ── ERD Diagram ────────────────────────────────────────────────────────────────
ERD_MERMAID_CODE = """
erDiagram
    trRMSample_Header {
        int intId PK
        varchar txtSampleNo UK
        int intStatusId
        int intCurrentStep
        varchar txtCreatedBy
        bit bitActive
    }
    trDocumentSample_Header {
        int intId PK
        int intSampleHeaderId FK
        varchar txtMaterialName
        varchar txtSupplierName
        int intStorageId FK
        varchar txtGMOStatus
        bit bitContainPHO
        varchar txtStatusOrganikNEW
        varchar txtEgDegContentNEW
        int intShelfLifeMonth
    }
    trDocumentSample_Allergen {
        int intId PK
        int intSampleHeaderId FK
        varchar txtAllergenCode
        bit bitIsChecked
    }
    trDocumentSample_BTP {
        int intId PK
        int intSampleHeaderId FK
        varchar txtBTPName
        decimal decPercentagePPM
    }
    trDocumentSample_Document {
        int intId PK
        int intSampleHeaderId FK
        varchar txtDocumentType
        varchar txtFileName
    }
    trSamplePurpose_Header {
        int intId PK
        int intSampleHeaderId FK
        varchar txtSamplePurposeType
        varchar txtItemCodeTrial
        varchar txtObjective
    }
    trRMEvaluation_Header {
        int intId PK
        int intSampleHeaderId FK
        decimal decOverallScore
        varchar txtMaterialComposition
    }
    trRMDisposition_Header {
        int intId PK
        int intSampleHeaderId FK
        varchar txtDecision
        bit bitApprovalRequired
        varchar txtApprovalStatus
    }

    trRMSample_Header ||--o| trDocumentSample_Header : "Step 1"
    trRMSample_Header ||--o{ trDocumentSample_Allergen : "Allergens"
    trRMSample_Header ||--o{ trDocumentSample_BTP : "BTP"
    trRMSample_Header ||--o{ trDocumentSample_Document : "Documents"
    trRMSample_Header ||--o| trSamplePurpose_Header : "Step 2"
    trRMSample_Header ||--o| trRMEvaluation_Header : "Step 3"
    trRMSample_Header ||--o| trRMDisposition_Header : "Step 4"
"""
ERD_PNG = os.path.join(SCREENSHOTS, 'new_rm_sample_erd_v2.png')

# ── Screenshot map ─────────────────────────────────────────────────────────────
SCREENSHOT_MAP = {
    'Gambar 1': FLOW_PNG,
    'Gambar 2': os.path.join(SCREENSHOTS, 'NewRMSample_Index_Page.png'),
    'Gambar 3': os.path.join(SCREENSHOTS, 'NewRMSample_Detail_Wizard.png'),
    'Gambar 4': os.path.join(SCREENSHOTS, 'Step1_Document_Sample.png'),
    'Gambar 5': os.path.join(SCREENSHOTS, 'Step2_Sample_Purpose.png'),
    'Gambar 6': os.path.join(SCREENSHOTS, 'Step3_RM_Evaluation.png'),
    'Gambar 7': os.path.join(SCREENSHOTS, 'Step4_Disposition.png'),
    'Gambar 8': ERD_PNG,
    # Additional detail screenshots for v2.0 accordion changes
    'Gambar 9': os.path.join(SCREENSHOTS, 'Halal_GMO_PHO_Accordion.png'),
    'Gambar 10': os.path.join(SCREENSHOTS, 'Storage_ShelfLife_Accordion.png'),
}

HEADER_BG    = '9BC2E6'   # light blue
BORDER_COLOR = '000000'
FONT_NAME    = 'Calibri'
FONT_SIZE    = 11


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 – Render mermaid diagrams
# ══════════════════════════════════════════════════════════════════════════════
def render_mermaid_png(mermaid_code, output_path, label):
    os.makedirs(SCREENSHOTS, exist_ok=True)
    encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('ascii')
    url = f'https://mermaid.ink/img/{encoded}?type=png&scale=3'
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = resp.read()
        with open(output_path, 'wb') as f:
            f.write(data)
        print(f'   [{label}] Saved {len(data):,} bytes -> {output_path}')
    except Exception as e:
        print(f'   [{label}] Warning: mermaid.ink failed ({e}). Using existing PNG if available.')


def render_mermaid():
    print('[1/4] Rendering mermaid diagrams...')
    render_mermaid_png(FLOW_MERMAID_CODE, FLOW_PNG, 'Flow')
    render_mermaid_png(ERD_MERMAID_CODE, ERD_PNG, 'ERD')


# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 – Pre-process markdown
# ══════════════════════════════════════════════════════════════════════════════
def preprocess_markdown():
    print('[2/4] Pre-processing markdown...')
    with open(MD_IN, 'r', encoding='utf-8') as f:
        text = f.read()

    # Replace ```mermaid ... ``` blocks with inline image
    def replace_mermaid_flow(m):
        if os.path.exists(FLOW_PNG):
            rel = os.path.relpath(FLOW_PNG, SCRIPT_DIR).replace('\\', '/')
            return f'\n![New RM Sample Business Flow Diagram v2.0]({rel})\n'
        return '*(Flow Diagram tidak tersedia)*\n'

    def replace_mermaid_erd(m):
        if os.path.exists(ERD_PNG):
            rel = os.path.relpath(ERD_PNG, SCRIPT_DIR).replace('\\', '/')
            return f'\n![New RM Sample ERD v2.0]({rel})\n'
        return '*(ERD Diagram tidak tersedia)*\n'

    # Replace ALL mermaid blocks — first is flowchart, second is erDiagram
    mermaid_blocks = list(re.finditer(r'```mermaid.*?```', text, flags=re.DOTALL))
    replacements = []
    for i, m in enumerate(mermaid_blocks):
        if i == 0:
            replacements.append((m.start(), m.end(), replace_mermaid_flow(m)))
        else:
            replacements.append((m.start(), m.end(), replace_mermaid_erd(m)))
    # Apply replacements in reverse order
    for start, end, repl in reversed(replacements):
        text = text[:start] + repl + text[end:]

    # Replace *(Gambar N: ...)* with inline image
    def replace_screenshot(m):
        label   = m.group(1).strip()
        caption = m.group(2).strip()
        for key, path in SCREENSHOT_MAP.items():
            if key in label and os.path.exists(path):
                rel = os.path.relpath(path, SCRIPT_DIR).replace('\\', '/')
                return f'\n![{caption}]({rel})\n'
        return f'\n*({label}: {caption} — screenshot belum tersedia)*\n'

    text = re.sub(r'\*\(?(Gambar\s*\d+)[:\s]*([^)]*)\)?\*', replace_screenshot, text)

    with open(MD_TMP, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f'   Written: {MD_TMP}')


# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 – Run pandoc
# ══════════════════════════════════════════════════════════════════════════════
def run_pandoc():
    print('[3/4] Running pandoc...')
    cmd = [
        'pandoc', MD_TMP,
        '-o', DOCX_OUT,
        '--from=markdown+pipe_tables',
        f'--resource-path={SCRIPT_DIR}',
        '--toc',
        '--toc-depth=3',
    ]
    if os.path.exists(REF_DOCX):
        cmd += [f'--reference-doc={REF_DOCX}']
        print(f'   Using reference doc: {REF_DOCX}')
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print('   pandoc stderr:', result.stderr)
        raise RuntimeError('pandoc failed')
    print(f'   Generated: {DOCX_OUT}')


# ══════════════════════════════════════════════════════════════════════════════
# STEP 4 – Post-process DOCX
# ══════════════════════════════════════════════════════════════════════════════
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_cell_bg_color(cell, color_str):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_str)
    cell._tc.get_or_add_tcPr().append(shading)


def apply_font(run, size_pt, bold=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    el = rPr.find(qn('w:rFonts'))
    if el is None:
        el = OxmlElement('w:rFonts')
        rPr.insert(0, el)
    el.set(qn('w:ascii'),  FONT_NAME)
    el.set(qn('w:hAnsi'), FONT_NAME)
    el.set(qn('w:cs'),    FONT_NAME)
    if bold is not None:
        run.bold = bold


def postprocess_docx():
    print('[4/4] Post-processing DOCX (borders, colours, fonts)...')
    doc = Document(DOCX_OUT)
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE)

    for para in doc.paragraphs:
        for run in para.runs:
            apply_font(run, FONT_SIZE)

    for table in doc.tables:
        try:
            table.style = 'Table Grid'
        except Exception:
            pass
        for row_idx, row in enumerate(table.rows):
            is_header = (row_idx == 0)
            for cell in row.cells:
                set_cell_border(
                    cell,
                    top={"sz": 4, "val": "single", "color": "000000"},
                    bottom={"sz": 4, "val": "single", "color": "000000"},
                    start={"sz": 4, "val": "single", "color": "000000"},
                    end={"sz": 4, "val": "single", "color": "000000"}
                )
                if is_header:
                    set_cell_bg_color(cell, HEADER_BG)
                for para in cell.paragraphs:
                    for run in para.runs:
                        apply_font(run, size_pt=10, bold=True if is_header else None)

    doc.save(DOCX_OUT)
    print(f'   Saved styled DOCX -> {DOCX_OUT}')


# ══════════════════════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print('=' * 60)
    print('  Building FSD New RM Sample v2.0 DOCX')
    print('=' * 60)
    render_mermaid()
    preprocess_markdown()
    run_pandoc()
    postprocess_docx()
    if os.path.exists(MD_TMP):
        os.remove(MD_TMP)
        print(f'   Cleaned up temp: {MD_TMP}')
    print()
    print('Done! Output file:', DOCX_OUT)
