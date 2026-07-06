"""
build_fsd_itemcodetrial_v3.py
=============================
Pipeline untuk membangun FSD Item Code Trial v3.0 format DOCX.
Menggunakan kroki.io untuk render diagram Mermaid.

Run:
    py build_fsd_itemcodetrial_v3.py
"""

import sys
import io
import base64
import os
import re
import subprocess
import urllib.request
import urllib.error
import shutil
import zlib

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# PATH CONFIGURATION
# ─────────────────────────────────────────────
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
MD_SRC      = os.path.join(SCRIPT_DIR, 'FSD_ItemCodeTrial_v3.0.md')
MD_TMP      = os.path.join(SCRIPT_DIR, '_tmp_itemcodetrial_v3_processed.md')
DOCX_OUT    = os.path.join(SCRIPT_DIR, 'FSD_ItemCodeTrial_v3.0.docx')
REF_DOCX    = os.path.join(SCRIPT_DIR, '..', 'reference.docx')
SCREENSHOTS = os.path.join(SCRIPT_DIR, 'screenshots')

os.makedirs(SCREENSHOTS, exist_ok=True)

# Diagram output files
FLOW_MAIN_PNG    = os.path.join(SCREENSHOTS, 'flow_main_v3.png')
FLOW_APPROVAL_PNG = os.path.join(SCREENSHOTS, 'flow_approval_v3.png')
ERD_PNG          = os.path.join(SCREENSHOTS, 'erd_v3.png')

# Style config
HEADER_BG    = 'D9EAD3'    # light green header
BORDER_COLOR = '000000'    # black border
FONT_NAME    = 'Calibri'
FONT_SIZE_BODY  = 11
FONT_SIZE_TABLE = 9


# ─────────────────────────────────────────────
# RENDER MERMAID VIA KROKI.IO
# ─────────────────────────────────────────────
def render_kroki(mermaid_code: str, output_path: str, label: str) -> bool:
    """Render Mermaid diagram ke PNG via kroki.io API (compressed)."""
    try:
        compressed = zlib.compress(mermaid_code.strip().encode('utf-8'), 9)
        b64 = base64.urlsafe_b64encode(compressed).decode('ascii')
        url = f'https://kroki.io/mermaid/png/{b64}'
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = resp.read()
        with open(output_path, 'wb') as f:
            f.write(data)
        print(f'   [{label}] OK -> {len(data):,} bytes -> {os.path.basename(output_path)}')
        return True
    except Exception as e:
        print(f'   [{label}] GAGAL: {e}')
        return False


# ─────────────────────────────────────────────
# STEP 1: PREPROCESS MARKDOWN
# ─────────────────────────────────────────────
def step1_preprocess():
    """Baca MD, render semua mermaid block, ganti dengan image link."""
    print('\n[STEP 1] Pre-processing Markdown dan render Mermaid diagrams...')

    with open(MD_SRC, 'r', encoding='utf-8') as f:
        text = f.read()

    # Find all mermaid blocks
    blocks = list(re.finditer(r'```mermaid\s*\n(.*?)```', text, flags=re.DOTALL))
    print(f'   Ditemukan {len(blocks)} mermaid block.')

    diagram_files = []
    for i, match in enumerate(blocks):
        code = match.group(1).strip()

        # Classify which diagram this is
        if 'XXSHP_INV_MASTER_ITEM_STG' in code and 'erDiagram' in code:
            png_path = ERD_PNG
            label = f'ERD-{i}'
        elif 'K2 BPM' in code or 'ApproverAct' in code:
            png_path = FLOW_APPROVAL_PNG
            label = f'Flow-Approval-{i}'
        else:
            png_path = os.path.join(SCREENSHOTS, f'diagram_{i}.png')
            label = f'Diagram-{i}'

        success = render_kroki(code, png_path, label)
        diagram_files.append((png_path, success))

    # Replace mermaid blocks with image references
    block_iter = iter(enumerate(diagram_files))

    def replace_mermaid(m):
        try:
            i, (png_path, success) = next(block_iter)
        except StopIteration:
            return m.group(0)

        rel_path = os.path.relpath(png_path, SCRIPT_DIR).replace('\\', '/')
        if success and os.path.exists(png_path):
            return f'\n![Diagram]({rel_path})\n'
        return '\n> *[Diagram render gagal - lihat FSD_ItemCodeTrial_v3.0.md untuk versi Mermaid]*\n'

    text = re.sub(r'```mermaid\s*\n(.*?)```', replace_mermaid, text, flags=re.DOTALL)

    with open(MD_TMP, 'w', encoding='utf-8') as f:
        f.write(text)

    print(f'   Markdown temp tersimpan: {os.path.basename(MD_TMP)}')


# ─────────────────────────────────────────────
# STEP 2: PANDOC CONVERSION
# ─────────────────────────────────────────────
def step2_pandoc():
    """Jalankan Pandoc untuk konversi MD -> DOCX."""
    print('\n[STEP 2] Menjalankan Pandoc (MD -> DOCX)...')

    cmd = [
        'pandoc', MD_TMP,
        '-o', DOCX_OUT,
        '--from=markdown+pipe_tables',
        f'--resource-path={SCRIPT_DIR};{SCREENSHOTS}',
        '--toc',
        '--toc-depth=3',
    ]

    if os.path.exists(REF_DOCX):
        cmd += [f'--reference-doc={REF_DOCX}']
        print(f'   Menggunakan reference doc: {os.path.basename(REF_DOCX)}')
    else:
        print(f'   [INFO] Reference doc tidak ditemukan, menggunakan style default.')

    result = subprocess.run(cmd, capture_output=True, text=True, cwd=SCRIPT_DIR)

    if result.returncode != 0:
        print(f'   [ERROR] Pandoc gagal:\n{result.stderr[:800]}')
        raise RuntimeError('Pandoc execution error')

    print(f'   DOCX berhasil dibuat: {os.path.basename(DOCX_OUT)}')


# ─────────────────────────────────────────────
# STEP 3: POSTPROCESS DOCX
# ─────────────────────────────────────────────
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = f'w:{edge}'
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for k, v in kwargs[edge].items():
                element.set(qn(f'w:{k}'), str(v))


def set_cell_bg(cell, color: str):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def apply_font(run, size_pt, bold=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    el = rPr.find(qn('w:rFonts'))
    if el is None:
        el = OxmlElement('w:rFonts')
        rPr.insert(0, el)
    el.set(qn('w:ascii'), FONT_NAME)
    el.set(qn('w:hAnsi'), FONT_NAME)
    el.set(qn('w:cs'), FONT_NAME)
    if bold is not None:
        run.bold = bold


def step3_postprocess():
    """Post-process DOCX: font, tabel borders, image scaling."""
    print('\n[STEP 3] Post-processing DOCX (Font, Tabel, Images)...')

    if not os.path.exists(DOCX_OUT):
        print('   [ERROR] DOCX tidak ditemukan! Skip postprocess.')
        return

    doc = Document(DOCX_OUT)

    # Set default font
    doc.styles['Normal'].font.name = FONT_NAME
    doc.styles['Normal'].font.size = Pt(FONT_SIZE_BODY)

    for para in doc.paragraphs:
        for run in para.runs:
            apply_font(run, FONT_SIZE_BODY)

    # Format tables
    border = {"sz": "8", "val": "single", "color": BORDER_COLOR}
    bdr_all = {
        'top': border, 'bottom': border,
        'start': border, 'end': border,
        'insideH': border, 'insideV': border
    }

    for tbl in doc.tables:
        try:
            tbl.style = 'Table Grid'
        except Exception:
            pass

        for row_idx, row in enumerate(tbl.rows):
            is_header = (row_idx == 0)
            for cell in row.cells:
                set_cell_border(cell, **bdr_all)
                if is_header:
                    set_cell_bg(cell, HEADER_BG)
                for para in cell.paragraphs:
                    for run in para.runs:
                        apply_font(run, FONT_SIZE_TABLE, bold=True if is_header else None)

    # Scale images to max width
    MAX_WIDTH = Cm(15)
    NS_DRAW = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    for para in doc.paragraphs:
        for run in para.runs:
            for inline in run._r.findall(f'.//{{{NS_DRAW}}}inline'):
                extent = inline.find(f'{{{NS_DRAW}}}extent')
                if extent is not None:
                    cx = int(extent.get('cx', 0))
                    if cx > MAX_WIDTH.emu:
                        ratio = MAX_WIDTH.emu / cx
                        cy = int(int(extent.get('cy', 0)) * ratio)
                        extent.set('cx', str(MAX_WIDTH.emu))
                        extent.set('cy', str(cy))

    doc.save(DOCX_OUT)
    print(f'   DOCX final tersimpan: {os.path.basename(DOCX_OUT)}')


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
if __name__ == '__main__':
    print('=' * 65)
    print('  BUILD FSD: Item Code Trial v3.0')
    print('  Output:', os.path.basename(DOCX_OUT))
    print('=' * 65)

    step1_preprocess()
    step2_pandoc()
    step3_postprocess()

    # Cleanup temp
    if os.path.exists(MD_TMP):
        os.remove(MD_TMP)
        print(f'\n   Temp file dibersihkan: {os.path.basename(MD_TMP)}')

    print('\n' + '=' * 65)
    print(f'  SELESAI: {DOCX_OUT}')
    print('=' * 65)
