"""
build_fsd_item_registration.py
Pipeline to build FSD Item Registration DOCX:
  1. Render mermaid diagrams (Flow Trial, Flow Prod, ERD) → PNG via mermaid.ink
  2. Pre-process markdown: replace ```mermaid blocks and *(Gambar N: ...)* placeholders
  3. Run pandoc with reference.docx for styles (with TOC)
  4. Post-process DOCX: Arial font + light-green header rows + dark-green borders

Run: py build_fsd_item_registration.py
"""

import base64
import os
import re
import subprocess
import urllib.request

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────────
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
FSD_DIR     = os.path.dirname(SCRIPT_DIR)  # parent = Document/FSD/

MD_IN       = os.path.join(SCRIPT_DIR, 'FSD_Item_Registration_v1.0.md')
MD_TMP      = os.path.join(SCRIPT_DIR, '_tmp_fsd_item_reg_processed.md')
DOCX_OUT    = os.path.join(SCRIPT_DIR, 'FSD - Item Registration v1.0.docx')
REF_DOCX    = os.path.join(FSD_DIR, 'reference.docx')
SCREENSHOTS = os.path.join(SCRIPT_DIR, 'screenshots')

# ── Styling constants ──────────────────────────────────────────────────────────
HEADER_BG    = 'C6EFCE'   # light green (same as existing FSD docs)
BORDER_COLOR = '538135'   # dark green border
FONT_NAME    = 'Arial'
FONT_SIZE    = 11

# ══════════════════════════════════════════════════════════════════════════════
# Mermaid Diagrams
# ══════════════════════════════════════════════════════════════════════════════

# A. Alur Item Trial – Horizontal Swimlane
MERMAID_TRIAL = """
flowchart LR
    START(["START"])

    subgraph INITIATOR ["👤 INITIATOR"]
        direction LR
        A1["Step 1: Buat Item Trial\\nPilih Template (RM TRIAL / PM TRIAL)"]
        A2["Step 2: Isi Data\\nItem Code, LOB, Packing & Pallet"]
        A3["Submit\\nItem Trial"]
    end

    subgraph SYSTEM ["⚙️ SYSTEM"]
        direction LR
        S1["Auto-fill:\\nItem Type, UOM, Inventory Info"]
        S2["Validasi Field Wajib\\n& Auto-Save Draft"]
        S3["Notifikasi ke\\nApprover"]
        S4["Kirim ke Oracle ERP\\n(Tabel Staging)"]
    end

    subgraph APPROVER ["✅ APPROVER"]
        direction LR
        AP1{"Telaah\\nItem Trial"}
        AP2["Reject +\\nCatatan"]
    end

    END1(["Item Trial\\nSUBMITTED TO ORACLE"])

    START --> A1
    A1 --> S1
    S1 --> A2
    A2 --> A3
    A3 --> S2
    S2 --> S3
    S3 --> AP1
    AP1 -->|Approved| S4
    AP1 -->|Rejected| AP2
    AP2 --> A2
    S4 --> END1

    style START  fill:#4CAF50,color:#fff,stroke:#388E3C
    style END1   fill:#1565C0,color:#fff,stroke:#0D47A1
    style AP2    fill:#F44336,color:#fff,stroke:#B71C1C
    style S3     fill:#FF9800,color:#fff,stroke:#E65100
    style AP1    fill:#FFF9C4,color:#333,stroke:#F9A825
    style INITIATOR fill:#E8F5E9,stroke:#388E3C,color:#1B5E20
    style SYSTEM    fill:#E3F2FD,stroke:#1565C0,color:#0D47A1
    style APPROVER  fill:#FFF3E0,stroke:#E65100,color:#BF360C
"""

# B. Alur Item Production – Horizontal Swimlane
MERMAID_PROD = """
flowchart LR
    START(["START\\n(Trial / Non-Trial Mode)"])

    subgraph INITIATOR ["👤 INITIATOR"]
        direction LR
        A1["Step 3: Buat Item Production\\nPilih Mode & Template (RM / PM)"]
        A2["Step 4: Isi Data Utama\\nSupplier, Compliance, Barcode/MD/Halal"]
        A3["Submit\\nItem Production"]
    end

    subgraph SYSTEM ["⚙️ SYSTEM"]
        direction LR
        S1["Aktifkan Field:\\nHalal (RM) / Forestry (PM)"]
        S2["Validasi Field Wajib\\n& Auto-Save Draft"]
        S3["Notifikasi ke\\nApprover"]
        S4["Kirim ke Oracle ERP\\n(Tabel Staging)"]
    end

    subgraph APPROVER ["✅ APPROVER"]
        direction LR
        AP1{"Telaah\\nItem Production"}
        AP2["Reject +\\nCatatan"]
    end

    END2(["Item Production\\nSUBMITTED TO ORACLE"])

    START --> A1
    A1 --> S1
    S1 --> A2
    A2 --> A3
    A3 --> S2
    S2 --> S3
    S3 --> AP1
    AP1 -->|Approved| S4
    AP1 -->|Rejected| AP2
    AP2 --> A2
    S4 --> END2

    style START  fill:#4CAF50,color:#fff,stroke:#388E3C
    style END2   fill:#1565C0,color:#fff,stroke:#0D47A1
    style AP2    fill:#F44336,color:#fff,stroke:#B71C1C
    style S3     fill:#FF9800,color:#fff,stroke:#E65100
    style AP1    fill:#FFF9C4,color:#333,stroke:#F9A825
    style INITIATOR fill:#E8F5E9,stroke:#388E3C,color:#1B5E20
    style SYSTEM    fill:#E3F2FD,stroke:#1565C0,color:#0D47A1
    style APPROVER  fill:#FFF3E0,stroke:#E65100,color:#BF360C
"""

# C. ERD Diagram
MERMAID_ERD = """
erDiagram
    XXSHP_INV_MASTER_ITEM_STG {
        NUMBER      INTERFACE_ID        PK
        VARCHAR2    PROCESS_FLAG
        VARCHAR2    ITEM_TEMPLATE
        VARCHAR2    SEGMENT1
        VARCHAR2    DESCRIPTION
        VARCHAR2    PRIMARY_UOM_CODE
        VARCHAR2    ITEM_TYPE
        VARCHAR2    LOB
        VARCHAR2    PROCESS_MODE
        NUMBER      VENDOR_ID
        NUMBER      VENDOR_SITE_ID
        VARCHAR2    SUPPLIER_NAME
        VARCHAR2    PRINCIPAL_NAME
        VARCHAR2    HALAL_COUNTRY
        VARCHAR2    CHALAL_NO
        DATE        CHALAL_EXPIRED_DATE
        VARCHAR2    HALAL_LOGO
        VARCHAR2    HALAL_BODY
        VARCHAR2    AKASIA_NUM
        VARCHAR2    CMD_NO
        DATE        CMD_VALID
        VARCHAR2    FORESTRY_CERT
        VARCHAR2    DOC_STATUS
        VARCHAR2    CREATED_BY
        DATE        CREATION_DATE
    }
    XXSHP_INV_MASTER_ITEM_STG_BARCODE {
        NUMBER      ID                  PK
        NUMBER      INTERFACE_ID        FK
        VARCHAR2    BARCODE
        VARCHAR2    PRODUCT_NAME
        DATE        CREATION_DATE
    }
    XXSHP_INV_MASTER_ITEM_STG_MD {
        NUMBER      ID                  PK
        NUMBER      INTERFACE_ID        FK
        VARCHAR2    MD_NUMBER
        VARCHAR2    MD_PRODUCT_NAME
        VARCHAR2    MD_PRODUCTION_SITE
        DATE        MD_EXPIRY_DATE
        DATE        CREATION_DATE
    }
    XXSHP_INV_MASTER_ITEM_STG_HALAL {
        NUMBER      ID                  PK
        NUMBER      INTERFACE_ID        FK
        VARCHAR2    HALAL_NUMBER
        VARCHAR2    HALAL_PRODUCT_NAME
        DATE        HALAL_EXPIRY_DATE
        DATE        CREATION_DATE
    }
    MTL_ITEM_TEMPLATES {
        NUMBER      TEMPLATE_ID         PK
        VARCHAR2    TEMPLATE_NAME
    }
    MTL_SYSTEM_ITEMS_B {
        NUMBER      INVENTORY_ITEM_ID   PK
        VARCHAR2    SEGMENT1
        VARCHAR2    DESCRIPTION
        VARCHAR2    PRIMARY_UOM_CODE
    }
    AP_SUPPLIERS {
        NUMBER      VENDOR_ID           PK
        VARCHAR2    VENDOR_NAME
    }
    AP_SUPPLIER_SITES_ALL {
        NUMBER      VENDOR_SITE_ID      PK
        NUMBER      VENDOR_ID           FK
        VARCHAR2    VENDOR_SITE_CODE
    }

    XXSHP_INV_MASTER_ITEM_STG ||--o{ XXSHP_INV_MASTER_ITEM_STG_BARCODE : "has Barcodes"
    XXSHP_INV_MASTER_ITEM_STG ||--o{ XXSHP_INV_MASTER_ITEM_STG_MD      : "has MD Numbers"
    XXSHP_INV_MASTER_ITEM_STG ||--o{ XXSHP_INV_MASTER_ITEM_STG_HALAL   : "has Halal Numbers"
    MTL_ITEM_TEMPLATES         ||--o{ XXSHP_INV_MASTER_ITEM_STG         : "templates"
    MTL_SYSTEM_ITEMS_B         ||--o{ XXSHP_INV_MASTER_ITEM_STG         : "item reference"
    AP_SUPPLIERS               ||--o{ XXSHP_INV_MASTER_ITEM_STG         : "supplier"
    AP_SUPPLIERS               ||--o{ AP_SUPPLIER_SITES_ALL             : "has sites"
    AP_SUPPLIER_SITES_ALL      ||--o{ XXSHP_INV_MASTER_ITEM_STG         : "principal"
"""

# ── PNG paths ──────────────────────────────────────────────────────────────────
TRIAL_PNG = os.path.join(SCREENSHOTS, 'item_reg_flow_trial.png')
PROD_PNG  = os.path.join(SCREENSHOTS, 'item_reg_flow_prod.png')
ERD_PNG   = os.path.join(SCREENSHOTS, 'item_reg_erd.png')

# ── Screenshot placeholder map (for *(Gambar N: ...)*) ────────────────────────
SCREENSHOT_MAP = {
    'Gambar 1': TRIAL_PNG,
    'Gambar 2': PROD_PNG,
    'Gambar 3': os.path.join(SCREENSHOTS, 'ss_trial_index.png'),
    'Gambar 4': os.path.join(SCREENSHOTS, 'ss_trial_detail.png'),
    'Gambar 5': os.path.join(SCREENSHOTS, 'ss_prod_index.png'),
    'Gambar 6': os.path.join(SCREENSHOTS, 'ss_prod_detail.png'),
    'Gambar 7': ERD_PNG,
}


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 – Render mermaid diagrams to PNG via mermaid.ink
# ══════════════════════════════════════════════════════════════════════════════
def render_mermaid_png(mermaid_code, output_path, label):
    os.makedirs(SCREENSHOTS, exist_ok=True)
    encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('ascii')
    url = f'https://mermaid.ink/img/{encoded}?type=png'
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = resp.read()
        with open(output_path, 'wb') as f:
            f.write(data)
        print(f'   [{label}] Saved {len(data):,} bytes -> {output_path}')
    except Exception as e:
        print(f'   [{label}] Warning: mermaid.ink failed ({e}). Using existing PNG if available.')


def render_mermaid():
    print('[1/4] Rendering mermaid diagrams...')
    render_mermaid_png(MERMAID_TRIAL, TRIAL_PNG, 'Flow Trial')
    render_mermaid_png(MERMAID_PROD,  PROD_PNG,  'Flow Prod')
    render_mermaid_png(MERMAID_ERD,   ERD_PNG,   'ERD')


# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 – Pre-process markdown
# ══════════════════════════════════════════════════════════════════════════════
def preprocess_markdown():
    print('[2/4] Pre-processing markdown...')
    with open(MD_IN, 'r', encoding='utf-8') as f:
        text = f.read()

    # Replace ```mermaid ... ``` blocks (not used in this MD, but kept for safety)
    def replace_mermaid_block(m):
        return '*(Diagram tersedia sebagai gambar di bawah)*\n'
    text = re.sub(r'```mermaid.*?```', replace_mermaid_block, text, flags=re.DOTALL)

    # Replace *(Gambar N: ...)* with inline image reference
    def replace_screenshot(m):
        label   = m.group(1).strip()   # e.g. "Gambar 1"
        caption = m.group(2).strip()   # e.g. "Flowchart Alur Item Trial"
        for key, path in SCREENSHOT_MAP.items():
            if key in label and os.path.exists(path):
                rel = os.path.relpath(path, SCRIPT_DIR).replace('\\', '/')
                return f'\n![{caption}]({rel})\n'
        return m.group(0)   # keep original if image not found

    text = re.sub(r'\*\(?(Gambar\s*\d+)[:\s]*([^)]*)\)?\*', replace_screenshot, text)

    with open(MD_TMP, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f'   Written: {MD_TMP}')


# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 – Run pandoc
# ══════════════════════════════════════════════════════════════════════════════
def run_pandoc():
    print('[3/4] Running pandoc...')
    cmd = [
        'pandoc', MD_TMP,
        '-o', DOCX_OUT,
        '--from=markdown+pipe_tables',
        f'--resource-path={SCRIPT_DIR}',
        '--toc',
        '--toc-depth=3',
    ]
    if os.path.exists(REF_DOCX):
        cmd += [f'--reference-doc={REF_DOCX}']
        print(f'   Using reference: {REF_DOCX}')
    else:
        print('   Warning: reference.docx not found, using pandoc defaults')

    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print('   pandoc stderr:', result.stderr)
        raise RuntimeError('pandoc failed')
    print(f'   Generated: {DOCX_OUT}')


# ══════════════════════════════════════════════════════════════════════════════
# STEP 4 – Post-process DOCX (Arial font + green headers + borders)
# ══════════════════════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color, sz=8):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def apply_font(run, bold=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    rPr = run._r.get_or_add_rPr()
    el  = rPr.find(qn('w:rFonts'))
    if el is None:
        el = OxmlElement('w:rFonts')
        rPr.insert(0, el)
    el.set(qn('w:ascii'), FONT_NAME)
    el.set(qn('w:hAnsi'), FONT_NAME)
    el.set(qn('w:cs'),    FONT_NAME)
    if bold is not None:
        run.bold = bold


def postprocess_docx():
    print('[4/4] Post-processing DOCX (Arial + green headers + borders)...')
    doc = Document(DOCX_OUT)

    # Apply Arial to all body paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            apply_font(run)

    # Process all tables
    for table in doc.tables:
        try:
            table.style = doc.styles['Table Grid']
        except Exception:
            pass
        for ri, row in enumerate(table.rows):
            is_header = (ri == 0)
            for cell in row.cells:
                set_cell_borders(cell, BORDER_COLOR)
                if is_header:
                    set_cell_bg(cell, HEADER_BG)
                for para in cell.paragraphs:
                    for run in para.runs:
                        apply_font(run, bold=True if is_header else None)

    doc.save(DOCX_OUT)
    print(f'   Saved: {DOCX_OUT}')


# ══════════════════════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    render_mermaid()
    preprocess_markdown()
    run_pandoc()
    postprocess_docx()

    # Cleanup temp
    if os.path.exists(MD_TMP):
        os.remove(MD_TMP)

    print(f'\nDone! Open: {DOCX_OUT}')
