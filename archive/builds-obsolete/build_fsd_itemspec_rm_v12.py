"""
build_fsd_itemspec_rm.py
=========================
Pipeline untuk membangun FSD Item Spec RM v1.2 format DOCX.
Menggunakan kroki.io untuk merender diagram Mermaid menjadi PNG,
lalu Pandoc untuk konversi MD -> DOCX, dan python-docx untuk post-processing.

Run:
    py build_fsd_itemspec_rm.py

Requirements:
    pip install python-docx
    pandoc (https://pandoc.org/installing.html) harus tersedia di PATH
"""
import sys
import io
# Force UTF-8 output to avoid cp1252 UnicodeEncodeError on Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

import base64
import os
import re
import subprocess
import sys
import urllib.request
import urllib.error
import zlib

ENGINE_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(ENGINE_ROOT, 'lib'))

from docx import Document
from fsd_cover_merge import (
    parse_md_cover_meta,
    strip_md_for_body,
    merge_cover_and_content,
    content_start_index,
    COVER_TABLE_COUNT,
)
from fsd_paths import COVER_TEMPLATE, LOGO_PATH
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ──────────────────────────────────────────────────
# PATHS
# ──────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MD_SRC     = os.path.join(SCRIPT_DIR, 'FSD_ItemSpec_RM_v1.2.md')
MD_TMP     = os.path.join(SCRIPT_DIR, '_tmp_itemspec_rm_v12_processed.md')
DOCX_CONTENT_TMP = os.path.join(SCRIPT_DIR, '_tmp_itemspec_rm_v12_content.docx')
DOCX_OUT   = os.path.join(SCRIPT_DIR, 'FSD_ItemSpec_RM_v1.2.docx')
COVER_META = {}

# Reference DOCX untuk styling (gunakan dari folder parent jika ada)
REF_DOCX_CANDIDATES = [
    os.path.join(SCRIPT_DIR, 'reference.docx'),
    os.path.join(os.path.dirname(SCRIPT_DIR), 'Item_Registration', 'reference.docx'),
    os.path.join(os.path.dirname(SCRIPT_DIR), 'reference.docx'),
]
REF_DOCX = next((p for p in REF_DOCX_CANDIDATES if os.path.exists(p)), None)

SCREENSHOTS = os.path.join(SCRIPT_DIR, 'screenshots')
os.makedirs(SCREENSHOTS, exist_ok=True)

# Output diagram paths
FLOW_MAIN_PNG     = os.path.join(SCREENSHOTS, 'itemspec_rm_flow_main.png')
FLOW_APPROVAL_PNG = os.path.join(SCREENSHOTS, 'itemspec_rm_flow_approval.png')
ERD_PNG           = os.path.join(SCREENSHOTS, 'itemspec_rm_erd.png')

# ──────────────────────────────────────────────────
# STYLE CONSTANTS (sesuai standar FSD IDC System)
# ──────────────────────────────────────────────────
HEADER_BG       = 'D9EAD3'   # light green
BORDER_COLOR    = '000000'   # black
FONT_NAME       = 'Calibri'
FONT_SIZE_BODY  = 11
FONT_SIZE_TABLE = 9


# ──────────────────────────────────────────────────
# STEP 1: Render Mermaid diagrams via Kroki.io
# ──────────────────────────────────────────────────

def render_kroki(mermaid_code: str, output_path: str, label: str) -> bool:
    """Render Mermaid code to PNG via Kroki.io API."""
    try:
        compressed = zlib.compress(mermaid_code.strip().encode('utf-8'), 9)
        b64 = base64.urlsafe_b64encode(compressed).decode('ascii')
        url = f'https://kroki.io/mermaid/png/{b64}'
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=45) as resp:
            data = resp.read()
        with open(output_path, 'wb') as f:
            f.write(data)
        print(f'   [{label}] [OK] {len(data):,} bytes -> {os.path.basename(output_path)}')
        return True
    except Exception as e:
        print(f'   [{label}] [X] Kroki gagal: {e}')
        return False


def step1_render_diagrams(md_content: str):
    """Extract all mermaid blocks and render each to PNG."""
    print('\n[STEP 1] Render diagram Mermaid via Kroki.io...')

    blocks = list(re.finditer(r'```mermaid\s*\n(.*?)```', md_content, flags=re.DOTALL))
    print(f'   Ditemukan {len(blocks)} blok Mermaid.')

    if not blocks:
        print('   [WARN] Tidak ada blok mermaid – diagram dilewati.')
        return

    for i, match in enumerate(blocks):
        code = match.group(1).strip()

        if 'flowchart LR' in code and 'subgraph PDV' in code and 'subgraph QAQS' in code:
            render_kroki(code, FLOW_MAIN_PNG, 'Flow-Main-Swimlane')
        elif 'flowchart LR' in code and 'Creator' in code and 'WaitApproval' in code:
            render_kroki(code, FLOW_APPROVAL_PNG, 'Flow-Approval')
        elif 'erDiagram' in code and 'TrItemSpecRM_Header' in code:
            render_kroki(code, ERD_PNG, 'ERD-Main')
        else:
            generic = os.path.join(SCREENSHOTS, f'itemspec_rm_diagram_{i + 1}.png')
            render_kroki(code, generic, f'Diagram-{i + 1}')


# ──────────────────────────────────────────────────
# STEP 2: Preprocess Markdown
# ──────────────────────────────────────────────────

def _rel(path: str) -> str:
    return os.path.relpath(path, SCRIPT_DIR).replace('\\', '/')


def step2_preprocess_markdown():
    global COVER_META
    print('\n[STEP 2] Pre-processing Markdown...')
    with open(MD_SRC, 'r', encoding='utf-8') as f:
        raw = f.read()

    COVER_META = parse_md_cover_meta(raw)
    text = strip_md_for_body(raw)

    # Render all diagrams first
    step1_render_diagrams(text)

    counter = [0]

    def replace_mermaid(m):
        code = m.group(1).strip()
        counter[0] += 1
        i = counter[0]

        if 'flowchart LR' in code and 'subgraph PDV' in code and 'subgraph QAQS' in code:
            png     = FLOW_MAIN_PNG
            caption = 'Business Flow Diagram Swimlane – Item Spec RM'
        elif 'flowchart LR' in code and 'Creator' in code and 'WaitApproval' in code:
            png     = FLOW_APPROVAL_PNG
            caption = 'Approval Flow – Item Spec RM'
        elif 'erDiagram' in code and 'TrItemSpecRM_Header' in code:
            png     = ERD_PNG
            caption = 'ERD – Entity Relationship Diagram Item Spec RM'
        else:
            png     = os.path.join(SCREENSHOTS, f'itemspec_rm_diagram_{i}.png')
            caption = f'Diagram {i}'

        if os.path.exists(png):
            return f'\n![{caption}]({_rel(png)})\n'
        else:
            return f'\n> *[{caption} – diagram tidak dapat di-render]*\n'

    text = re.sub(r'```mermaid\s*\n(.*?)```', replace_mermaid, text, flags=re.DOTALL)

    with open(MD_TMP, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f'   [OK] Preprocessed MD -> {os.path.basename(MD_TMP)}')


# ──────────────────────────────────────────────────
# STEP 3: Run Pandoc (MD → DOCX)
# ──────────────────────────────────────────────────

def step3_run_pandoc():
    print('\n[STEP 3] Menjalankan Pandoc (MD -> DOCX, isi bab saja)...')

    cmd = [
        'pandoc', MD_TMP,
        '-o', DOCX_CONTENT_TMP,
        '--from=markdown+pipe_tables',
        f'--resource-path={SCRIPT_DIR};{SCREENSHOTS}',
    ]

    if REF_DOCX:
        cmd += [f'--reference-doc={REF_DOCX}']
        print(f'   Menggunakan reference.docx: {REF_DOCX}')
    else:
        print('   [WARN] reference.docx tidak ditemukan – menggunakan style default pandoc.')

    result = subprocess.run(cmd, capture_output=True, text=True, cwd=SCRIPT_DIR)
    if result.returncode != 0:
        print('   Pandoc GAGAL!\n', result.stderr[:800])
        raise RuntimeError('pandoc execution error')
    print(f'   [OK] DOCX konten: {os.path.basename(DOCX_CONTENT_TMP)}')


def step3b_merge_cover():
    print('\n[STEP 3b] Gabung cover + revision template (Kalbe standard)...')
    merge_cover_and_content(
        DOCX_CONTENT_TMP,
        DOCX_OUT,
        COVER_META,
        template_path=COVER_TEMPLATE,
        logo_path=LOGO_PATH,
    )
    print(f'   [OK] {os.path.basename(DOCX_OUT)} (cover 2 halaman + konten)')


# ──────────────────────────────────────────────────
# STEP 4: Post-process DOCX
# ──────────────────────────────────────────────────

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = f'w:{edge}'
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for k, v in kwargs[edge].items():
                element.set(qn(f'w:{k}'), str(v))


def set_cell_bg(cell, color: str):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def apply_font(run, size_pt: int, bold=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    el = rPr.find(qn('w:rFonts'))
    if el is None:
        el = OxmlElement('w:rFonts')
        rPr.insert(0, el)
    el.set(qn('w:ascii'), FONT_NAME)
    el.set(qn('w:hAnsi'), FONT_NAME)
    el.set(qn('w:cs'), FONT_NAME)
    if bold is not None:
        run.bold = bold


def step4_postprocess_docx():
    print('\n[STEP 4] Post-processing DOCX (Fonts, Table Borders, Image Scale)...')
    doc = Document(DOCX_OUT)
    content_start = content_start_index(doc)

    # Global font style
    try:
        doc.styles['Normal'].font.name = FONT_NAME
        doc.styles['Normal'].font.size = Pt(FONT_SIZE_BODY)
    except Exception:
        pass

    # Apply font to content paragraphs only (skip cover)
    for i, para in enumerate(doc.paragraphs):
        if i < content_start:
            continue
        for run in para.runs:
            apply_font(run, FONT_SIZE_BODY)

    border_spec = {"sz": "8", "val": "single", "color": BORDER_COLOR}
    bdr_all = dict(
        top=border_spec, bottom=border_spec,
        start=border_spec, end=border_spec,
        insideH=border_spec, insideV=border_spec
    )

    # Style content tables only (skip cover + revision)
    for ti, table in enumerate(doc.tables):
        if ti < COVER_TABLE_COUNT:
            continue
        try:
            table.style = 'Table Grid'
        except Exception:
            pass

        for row_idx, row in enumerate(table.rows):
            is_header = (row_idx == 0)
            for cell in row.cells:
                set_cell_border(cell, **bdr_all)
                if is_header:
                    set_cell_bg(cell, HEADER_BG)
                for para in cell.paragraphs:
                    for run in para.runs:
                        apply_font(run, FONT_SIZE_TABLE, bold=True if is_header else None)

    # Scale images to max 15 cm width
    MAX_WIDTH = Cm(15)
    for para in doc.paragraphs:
        for run in para.runs:
            drawings = run._r.findall(
                './/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'
            )
            for drawing in drawings:
                extent = drawing.find(
                    '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'
                )
                if extent is not None:
                    cx = int(extent.get('cx', 0))
                    if cx > MAX_WIDTH.emu:
                        ratio = MAX_WIDTH.emu / cx
                        cy = int(int(extent.get('cy', 0)) * ratio)
                        extent.set('cx', str(MAX_WIDTH.emu))
                        extent.set('cy', str(cy))

    doc.save(DOCX_OUT)
    print(f'   [OK] DOCX final tersimpan: {os.path.basename(DOCX_OUT)}')


# ──────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────

if __name__ == '__main__':
    print('=' * 65)
    print('  BUILD START: FSD Item Spec RM v1.2')
    print('  Modul: Raw Material Specification – IDC System')
    print('=' * 65)

    if not os.path.exists(MD_SRC):
        print(f'\n  ERROR: Source MD tidak ditemukan: {MD_SRC}')
        exit(1)

    step2_preprocess_markdown()
    step3_run_pandoc()
    step3b_merge_cover()
    step4_postprocess_docx()

    # Cleanup temp files
    for tmp in (MD_TMP, DOCX_CONTENT_TMP):
        if os.path.exists(tmp):
            os.remove(tmp)
            print(f'\n   [OK] Temp file dihapus: {os.path.basename(tmp)}')

    print('\n' + '=' * 65)
    print(f'  SELESAI: {DOCX_OUT}\n  Perubahan v1.2: Flowchart Swimlane, Hapus Struktur Halaman, Kartu -> Card')
    print('=' * 65)
