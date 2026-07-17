"""Shared build helpers for FSD Generator Engine modules."""
import base64
import os
import re
import subprocess
import urllib.request
import zlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from fsd_cover_merge import content_start_index, COVER_TABLE_COUNT
from fsd_captions import caption_number_from_text
from fsd_paths import REFERENCE_DOCX

HEADER_BG = 'D9EAD3'
BORDER_COLOR = '000000'
FONT_NAME = 'Calibri'
FONT_SIZE_BODY = 11
FONT_SIZE_TABLE = 9
HEADING_FONT_SIZES = {
    'Heading 1': 20,
    'Heading 2': 18,
    'Heading 3': 15,
    'Heading 4': 13,
}
BUTTON_IMAGE_MAX_CM = 4.0
TABLE_CELL_MARGIN_DXA = 140
BUTTON_TABLE_COL_WIDTHS_CM = (5.0, 3.5, 3.5, 2.5, 5.5)  # Tampilan, Tombol, ID, Style, Fungsi


def render_kroki(mermaid_code: str, output_path: str, label: str) -> bool:
    return _render_kroki(mermaid_code, output_path, label, 'mermaid')


def render_kroki_plantuml(plantuml_code: str, output_path: str, label: str) -> bool:
    return _render_kroki(plantuml_code, output_path, label, 'plantuml')


def _render_kroki(code: str, output_path: str, label: str, engine: str) -> bool:
    try:
        compressed = zlib.compress(code.strip().encode('utf-8'), 9)
        b64 = base64.urlsafe_b64encode(compressed).decode('ascii')
        url = f'https://kroki.io/{engine}/png/{b64}'
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = resp.read()
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        with open(output_path, 'wb') as f:
            f.write(data)
        print(f'   [{label}] OK {len(data):,} bytes -> {os.path.basename(output_path)}')
        return True
    except Exception as e:
        print(f'   [{label}] FAIL: {e}')
        return False


def ensure_png_min_width(path: str, min_width: int = 2400) -> None:
    """Upscale PNG so ERD/diagram fills print page width when embedded."""
    if not path or not os.path.exists(path):
        return
    try:
        from PIL import Image
        with Image.open(path) as img:
            w, h = img.size
            if w >= min_width:
                return
            scale = min_width / float(w)
            new_size = (min_width, max(1, int(h * scale)))
            resized = img.resize(new_size, Image.Resampling.LANCZOS)
            resized.save(path)
            print(f'   [PNG upscale] {os.path.basename(path)} {w}x{h} -> {new_size[0]}x{new_size[1]}')
    except Exception as e:
        print(f'   [PNG upscale] skip {os.path.basename(path)}: {e}')


def is_swimlane_mermaid(code: str) -> bool:
    """True if Mermaid code uses 2+ subgraph lanes (cross-functional swimlane)."""
    return len(re.findall(r'^\s*subgraph\s+', code, flags=re.MULTILINE)) >= 2


def is_swimlane_plantuml(code: str) -> bool:
    """True if PlantUML activity diagram uses swimlane separators (|Role|)."""
    return len(re.findall(r'^\s*\|[^|]+\|\s*$', code, flags=re.MULTILINE)) >= 2


PLANTUML_SWIMLANE_SKINPARAM = """
skinparam SwimlaneBorderColor #000000
skinparam SwimlaneBorderThickness 3
skinparam SwimlaneTitleBackgroundColor #D9EAD3
skinparam SwimlaneTitleFontStyle bold
skinparam ActivityBackgroundColor #FFFFFF
skinparam ActivityBorderColor #000000
skinparam ActivityStartColor #C8E6C9
skinparam ActivityEndColor #B2DFDB
""".strip()


def inject_plantuml_swimlane_style(code: str) -> str:
    """Inject swimlane border/title skinparams after @startuml when missing."""
    if not is_swimlane_plantuml(code):
        return code
    if re.search(r'SwimlaneBorderColor', code, flags=re.I):
        return code
    code = re.sub(r'skinparam\s+partition\s*\{[^}]*\}\s*', '', code, flags=re.I | re.DOTALL)
    m = re.search(r'(@startuml[^\n]*\n)', code, flags=re.I)
    if not m:
        return code
    insert_at = m.end()
    return code[:insert_at] + PLANTUML_SWIMLANE_SKINPARAM + '\n' + code[insert_at:]


def run_pandoc(md_tmp: str, docx_out: str, resource_dirs: list[str], cwd: str):
    paths = ';'.join(resource_dirs)
    cmd = [
        'pandoc', md_tmp, '-o', docx_out,
        '--from=markdown+pipe_tables',
        f'--resource-path={paths}',
    ]
    if os.path.exists(REFERENCE_DOCX):
        cmd.append(f'--reference-doc={REFERENCE_DOCX}')
    result = subprocess.run(cmd, capture_output=True, text=True, cwd=cwd)
    if result.returncode != 0:
        raise RuntimeError(result.stderr[:800])


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = f'w:{edge}'
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for k, v in kwargs[edge].items():
                element.set(qn(f'w:{k}'), str(v))


def _set_cell_bg(cell, color: str):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def _apply_font(run, size_pt: int, bold=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    el = rPr.find(qn('w:rFonts'))
    if el is None:
        el = OxmlElement('w:rFonts')
        rPr.insert(0, el)
    el.set(qn('w:ascii'), FONT_NAME)
    el.set(qn('w:hAnsi'), FONT_NAME)
    el.set(qn('w:cs'), FONT_NAME)
    if bold is not None:
        run.bold = bold


def _set_cell_margins(cell, margin_dxa: int = TABLE_CELL_MARGIN_DXA):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcMar = OxmlElement('w:tcMar')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(margin_dxa))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _is_button_action_table(table) -> bool:
    if not table.rows:
        return False
    header = ' '.join(c.text for c in table.rows[0].cells)
    return 'Tampilan' in header and 'Tombol' in header


def _scale_inline_drawing(drawing, max_w, min_w=None, force_w=None):
    ns = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    extent = drawing.find(f'{ns}extent')
    if extent is None:
        return
    cx = int(extent.get('cx', 0))
    cy = int(extent.get('cy', 0))
    if cx <= 0:
        return
    target = None
    if force_w is not None:
        target = force_w.emu
    elif cx > max_w.emu:
        target = max_w.emu
    elif min_w is not None and cx < min_w.emu:
        target = min_w.emu
    if target is None:
        return
    ratio = target / cx
    extent.set('cx', str(int(target)))
    extent.set('cy', str(max(1, int(cy * ratio))))


def _insert_page_break_before(para):
    """Sisipkan page break tepat sebelum paragraf."""
    new_p = OxmlElement('w:p')
    new_r = OxmlElement('w:r')
    new_br = OxmlElement('w:br')
    new_br.set(qn('w:type'), 'page')
    new_r.append(new_br)
    new_p.append(new_r)
    para._element.addprevious(new_p)


_CHAPTER_HEADING_RE = re.compile(r'^(\d+)\.\s+')


def postprocess_page_breaks(doc, content_start: int):
    """Page break setelah Daftar Isi (sebelum bab 1) dan sebelum setiap bab utama (## 2., ## 3., …)."""
    for i, para in enumerate(doc.paragraphs):
        style = (para.style.name if para.style else '') or ''
        if not style.startswith('Heading'):
            continue
        text = (para.text or '').strip()
        m = _CHAPTER_HEADING_RE.match(text)
        if not m:
            continue
        if i < content_start:
            continue
        chapter = int(m.group(1))
        if i == content_start or chapter >= 2:
            _insert_page_break_before(para)


def postprocess_headings(doc, content_start: int):
    """Perbesar font heading bab/section agar lebih mudah dibaca."""
    for i, para in enumerate(doc.paragraphs):
        if i < content_start:
            continue
        style = (para.style.name if para.style else '') or ''
        size = HEADING_FONT_SIZES.get(style)
        if not size:
            continue
        for run in para.runs:
            _apply_font(run, size, bold=True)


def postprocess_captions(doc, content_start: int):
    for i, para in enumerate(doc.paragraphs):
        if i < content_start:
            continue
        raw = (para.text or '').strip()
        if not raw or not caption_number_from_text(raw):
            continue
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.italic = True
            _apply_font(run, FONT_SIZE_BODY)


def postprocess_docx(
    docx_path: str,
    max_width_cm: float = 15.0,
    erd_width_cm: float | None = None,
):
    doc = Document(docx_path)
    content_start = content_start_index(doc)
    try:
        doc.styles['Normal'].font.name = FONT_NAME
        doc.styles['Normal'].font.size = Pt(FONT_SIZE_BODY)
    except Exception:
        pass
    for i, para in enumerate(doc.paragraphs):
        if i < content_start:
            continue
        style = (para.style.name if para.style else '') or ''
        if style.startswith('Heading'):
            continue
        for run in para.runs:
            _apply_font(run, FONT_SIZE_BODY)
    border_spec = {"sz": "8", "val": "single", "color": BORDER_COLOR}
    bdr_all = dict(top=border_spec, bottom=border_spec, start=border_spec, end=border_spec,
                   insideH=border_spec, insideV=border_spec)
    for ti, table in enumerate(doc.tables):
        if ti < COVER_TABLE_COUNT:
            continue
        is_btn_table = _is_button_action_table(table)
        try:
            table.style = 'Table Grid'
        except Exception:
            pass
        for row_idx, row in enumerate(table.rows):
            is_header = row_idx == 0
            for cell in row.cells:
                _set_cell_border(cell, **bdr_all)
                _set_cell_margins(cell)
                if is_header:
                    _set_cell_bg(cell, HEADER_BG)
                for para in cell.paragraphs:
                    for run in para.runs:
                        _apply_font(run, FONT_SIZE_TABLE, bold=True if is_header else None)
                    if is_btn_table:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    btn_max = Cm(BUTTON_IMAGE_MAX_CM)
                    ns = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
                    for run in para.runs:
                        for drawing in run._r.findall(f'.//{ns}inline'):
                            _scale_inline_drawing(drawing, btn_max)
        if is_btn_table and len(BUTTON_TABLE_COL_WIDTHS_CM) == len(table.columns):
            for ci, w in enumerate(BUTTON_TABLE_COL_WIDTHS_CM):
                table.columns[ci].width = Cm(w)
    max_w = Cm(max_width_cm)
    erd_w = Cm(erd_width_cm) if erd_width_cm else None
    ns = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    paras = list(doc.paragraphs)
    for i, para in enumerate(paras):
        next_text = (paras[i + 1].text or '').upper() if i + 1 < len(paras) else ''
        para_text = (para.text or '').upper()
        is_erd = 'ERD' in para_text or 'ERD' in next_text
        for run in para.runs:
            for drawing in run._r.findall(f'.//{ns}inline'):
                doc_pr = None
                for el in drawing.iter():
                    if el.tag.endswith('docPr'):
                        doc_pr = el
                        break
                if doc_pr is not None:
                    name = (doc_pr.get('name') or '') + ' ' + (doc_pr.get('descr') or '')
                    if 'ERD' in name.upper():
                        is_erd = True
                if is_erd and erd_w is not None:
                    _scale_inline_drawing(drawing, max_w, force_w=erd_w)
                else:
                    _scale_inline_drawing(drawing, max_w)
    postprocess_page_breaks(doc, content_start)
    postprocess_headings(doc, content_start)
    postprocess_captions(doc, content_start)
    doc.save(docx_path)
