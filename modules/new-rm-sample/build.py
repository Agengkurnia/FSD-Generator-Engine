"""
Build FSD New RM Sample v3.1 → DOCX (Kalbe cover standard).

Run:
    py build.py
"""
import base64
import os
import re
import subprocess
import sys
import shutil
import zlib

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(SCRIPT_DIR)), 'lib'))

from fsd_cover_merge import (
    parse_md_cover_meta,
    strip_md_for_body,
    merge_cover_and_content,
    content_start_index,
    COVER_TABLE_COUNT,
)
from fsd_paths import REFERENCE_DOCX, TMP_DIR, engine_root_from_script

ENGINE_ROOT = engine_root_from_script(__file__)
MODULE_SLUG = 'new-rm-sample'
os.makedirs(TMP_DIR, exist_ok=True)
os.makedirs(os.path.join(SCRIPT_DIR, 'output'), exist_ok=True)

MD_SRC = os.path.join(SCRIPT_DIR, 'source', 'FSD_New_RM_Sample_v3.1.md')
MD_TMP = os.path.join(TMP_DIR, f'{MODULE_SLUG}_processed.md')
DOCX_CONTENT_TMP = os.path.join(TMP_DIR, f'{MODULE_SLUG}_content.docx')
DOCX_OUT = os.path.join(SCRIPT_DIR, 'output', 'FSD_New_RM_Sample_v3.1.docx')
SCREENSHOTS = os.path.join(SCRIPT_DIR, 'screenshots')
COVER_META = {}

HEADER_BG = 'D9EAD3'
BORDER_COLOR = '000000'
FONT_NAME = 'Calibri'
FONT_SIZE_BODY = 11
FONT_SIZE_TABLE = 9

os.makedirs(SCREENSHOTS, exist_ok=True)
ERD_WORKFLOW_PNG = os.path.join(SCREENSHOTS, 'erd_workflow_v3_1.png')
ERD_MASTER_PNG = os.path.join(SCREENSHOTS, 'erd_master_v3_1.png')
FLOW_MAIN_PNG = os.path.join(SCREENSHOTS, 'flow_main_v3_1.png')

UI_SCREENSHOT_KEYS = [
    'step1_header_wizard', 'step1_supplier_material', 'step1_pricing_packaging',
    'step1_storage_shelflife', 'step1_halal_gmo', 'step1_allergen', 'step1_btp_content',
    'step2_full_view', 'step2_lov_popup', 'step2_document_upload',
]


def render_kroki(mermaid_code, output_path, label):
    try:
        compressed = zlib.compress(mermaid_code.strip().encode('utf-8'), 9)
        b64 = base64.urlsafe_b64encode(compressed).decode('ascii')
        url = f'https://kroki.io/mermaid/png/{b64}'
        req = __import__('urllib.request').request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with __import__('urllib.request').request.urlopen(req, timeout=45) as resp:
            data = resp.read()
        with open(output_path, 'wb') as f:
            f.write(data)
        print(f'   [{label}] OK {len(data):,} bytes')
        return True
    except Exception as e:
        print(f'   [{label}] FAIL: {e}')
        return False


def extract_mermaids_and_render(md_content):
    print('\n[STEP 1] Render Mermaid (Kroki)...')
    for match in re.finditer(r'```mermaid\s*\n(.*?)```', md_content, flags=re.DOTALL):
        code = match.group(1).strip()
        if 'trRMSample_Header ||--o| trDocumentSample_Header' in code:
            render_kroki(code, ERD_WORKFLOW_PNG, 'ERD-Workflow')
        elif 'mSupplier' in code and 'trDocumentSample_Header }o--|| mSupplier' in code:
            render_kroki(code, ERD_MASTER_PNG, 'ERD-Master')
        elif 'flowchart ' in code and 'Start([' in code:
            render_kroki(code, FLOW_MAIN_PNG, 'Flow-Main')


def step2_copy_screenshots():
    """Salin screenshot opsional dari env FSD_SCREENSHOT_SRC jika ada."""
    ext = os.environ.get('FSD_SCREENSHOT_SRC', '').strip()
    if not ext or not os.path.isdir(ext):
        print('\n[STEP 2] Screenshot: pakai folder screenshots/ modul (set FSD_SCREENSHOT_SRC untuk salin eksternal)')
        return
    print(f'\n[STEP 2] Salin screenshot dari {ext}')
    for key in UI_SCREENSHOT_KEYS:
        for fname in os.listdir(ext):
            if key in fname and fname.endswith('.png'):
                shutil.copy2(os.path.join(ext, fname), os.path.join(SCREENSHOTS, f'{key}.png'))
                print(f'   OK {key}.png')
                break


def _rel(path):
    return os.path.relpath(path, SCRIPT_DIR).replace('\\', '/')


def step3_preprocess():
    global COVER_META
    print('\n[STEP 3] Pre-process Markdown...')
    with open(MD_SRC, 'r', encoding='utf-8') as f:
        raw = f.read()
    COVER_META = parse_md_cover_meta(raw)
    text = strip_md_for_body(raw)
    extract_mermaids_and_render(text)

    def replace_mermaid_block(m):
        code = m.group(1).strip()
        if 'trRMSample_Header ||--o| trDocumentSample_Header' in code:
            png, caption = ERD_WORKFLOW_PNG, 'ERD – Transaksi Workflow'
        elif 'mSupplier' in code and 'trDocumentSample_Header }o--|| mSupplier' in code:
            png, caption = ERD_MASTER_PNG, 'ERD – Relasi Tabel Master'
        elif 'flowchart LR' in code and 'Submit' in code:
            return '\n> *[Diagram: Logika Submit – refer ke Deskripsi Flow]*\n'
        elif 'flowchart ' in code and 'Start([' in code:
            png, caption = FLOW_MAIN_PNG, 'Business Flow Diagram'
        elif 'Item Code Trial' in code and 'Tampilkan LOV' in code:
            return '\n> *[Flowchart Auto Generate]*\n'
        else:
            return m.group(0)
        if os.path.exists(png):
            return f'\n![{caption}]({_rel(png)})\n'
        return f'\n*({caption} – diagram gagal di-render)*\n'

    text = re.sub(r'```mermaid\s*\n(.*?)```', replace_mermaid_block, text, flags=re.DOTALL)

    def replace_abs_img(m):
        full_path = m.group(2).replace('\\', '/')
        fname = os.path.basename(full_path)
        candidates = [os.path.join(SCREENSHOTS, fname)]
        for key in UI_SCREENSHOT_KEYS:
            if key in fname:
                candidates.append(os.path.join(SCREENSHOTS, f'{key}.png'))
        for candidate in candidates:
            if os.path.exists(candidate):
                return f'![{m.group(1)}]({_rel(candidate)})'
        if os.path.exists(full_path):
            return f'![{m.group(1)}]({full_path})'
        return f'*[Screenshot: {m.group(1)} – Image Missing]*'

    text = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', replace_abs_img, text)
    with open(MD_TMP, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f'   OK -> {os.path.basename(MD_TMP)}')


def step4_pandoc():
    print('\n[STEP 4] Pandoc...')
    cmd = [
        'pandoc', MD_TMP, '-o', DOCX_CONTENT_TMP,
        '--from=markdown+pipe_tables',
        f'--resource-path={SCRIPT_DIR};{SCREENSHOTS}',
    ]
    if os.path.exists(REFERENCE_DOCX):
        cmd.append(f'--reference-doc={REFERENCE_DOCX}')
    result = subprocess.run(cmd, capture_output=True, text=True, cwd=SCRIPT_DIR)
    if result.returncode != 0:
        raise RuntimeError(result.stderr[:500])


def step4b_merge_cover():
    print('\n[STEP 4b] Merge cover...')
    merge_cover_and_content(DOCX_CONTENT_TMP, DOCX_OUT, COVER_META)


def apply_font(run, size_pt, bold=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    el = rPr.find(qn('w:rFonts'))
    if el is None:
        el = OxmlElement('w:rFonts')
        rPr.insert(0, el)
    el.set(qn('w:ascii'), FONT_NAME)
    el.set(qn('w:hAnsi'), FONT_NAME)
    el.set(qn('w:cs'), FONT_NAME)
    if bold is not None:
        run.bold = bold


def step5_postprocess():
    print('\n[STEP 5] Post-process...')
    doc = Document(DOCX_OUT)
    content_start = content_start_index(doc)
    doc.styles['Normal'].font.name = FONT_NAME
    doc.styles['Normal'].font.size = Pt(FONT_SIZE_BODY)
    for i, para in enumerate(doc.paragraphs):
        if i < content_start:
            continue
        for run in para.runs:
            apply_font(run, FONT_SIZE_BODY)
    border = {"sz": "8", "val": "single", "color": BORDER_COLOR}
    bdr_all = dict(top=border, bottom=border, start=border, end=border, insideH=border, insideV=border)
    for ti, table in enumerate(doc.tables):
        if ti < COVER_TABLE_COUNT:
            continue
        try:
            table.style = 'Table Grid'
        except Exception:
            pass
        for row_idx, row in enumerate(table.rows):
            is_header = row_idx == 0
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is None:
                    tcBorders = OxmlElement('w:tcBorders')
                    tcPr.append(tcBorders)
                for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
                    element = tcBorders.find(qn(f'w:{edge}'))
                    if element is None:
                        element = OxmlElement(f'w:{edge}')
                        tcBorders.append(element)
                    for k, v in border.items():
                        element.set(qn(f'w:{k}'), str(v))
                if is_header:
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:fill'), HEADER_BG)
                    tcPr.append(shd)
                for para in cell.paragraphs:
                    for run in para.runs:
                        apply_font(run, FONT_SIZE_TABLE, bold=True if is_header else None)
    max_w = Cm(15)
    ns = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    for para in doc.paragraphs:
        for run in para.runs:
            for drawing in run._r.findall(f'.//{ns}inline'):
                extent = drawing.find(f'{ns}extent')
                if extent is not None:
                    cx = int(extent.get('cx', 0))
                    if cx > max_w.emu:
                        ratio = max_w.emu / cx
                        cy = int(int(extent.get('cy', 0)) * ratio)
                        extent.set('cx', str(max_w.emu))
                        extent.set('cy', str(cy))
    doc.save(DOCX_OUT)
    print(f'   OK -> {DOCX_OUT}')


if __name__ == '__main__':
    print('=' * 60)
    print('  BUILD: FSD New RM Sample v3.1')
    print('=' * 60)
    if not os.path.exists(MD_SRC):
        print(f'ERROR: {MD_SRC}')
        sys.exit(1)
    step2_copy_screenshots()
    step3_preprocess()
    step4_pandoc()
    step4b_merge_cover()
    step5_postprocess()
    for tmp in (MD_TMP, DOCX_CONTENT_TMP):
        if os.path.exists(tmp):
            os.remove(tmp)
    print(f'\nSELESAI: {DOCX_OUT}')
