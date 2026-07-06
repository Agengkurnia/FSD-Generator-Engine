"""
build_fsd_item_trial_formulation.py
======================================
Pipeline untuk membangun FSD Item Trial Formulation v1.0 format DOCX.
Menggunakan kroki.io untuk merender diagram Mermaid menjadi PNG.

Run:
    py build_fsd_item_trial_formulation.py

Requirements:
    pip install python-docx
    pandoc (https://pandoc.org/installing.html) harus tersedia di PATH
"""
import sys
import io
# Force UTF-8 output to avoid cp1252 UnicodeEncodeError on Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

import base64
import os
import re
import subprocess
import urllib.request
import urllib.error
import shutil
import zlib

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ──────────────────────────────────────────────────
# PATHS
# ──────────────────────────────────────────────────
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
MD_SRC      = os.path.join(SCRIPT_DIR, 'FSD_Item_Trial_Formulation_v1.0.md')
MD_TMP      = os.path.join(SCRIPT_DIR, '_tmp_itf_v10_processed.md')
DOCX_OUT    = os.path.join(SCRIPT_DIR, 'FSD_Item_Trial_Formulation_v1.0.docx')
SCREENSHOTS = os.path.join(SCRIPT_DIR, '..', 'screenshots')  # relative to brain dir

# Try to find reference docx from Item Registration FSD
REF_DOCX_CANDIDATES = [
    r"d:\Work\Source\RMSelection Hijrah\NewRMSelection\Document\FSD\Item_Registration\reference.docx",
    os.path.join(SCRIPT_DIR, 'reference.docx'),
]
REF_DOCX = next((p for p in REF_DOCX_CANDIDATES if os.path.exists(p)), None)

os.makedirs(SCREENSHOTS, exist_ok=True)

# Output diagram paths - screenshots are in parent_dir\screenshots
BRAIN_DIR       = os.path.dirname(SCRIPT_DIR)   # parent of FSD_Formulation
SCREENSHOTS_DIR = os.path.join(BRAIN_DIR, 'screenshots')
os.makedirs(SCREENSHOTS_DIR, exist_ok=True)

FLOW_BUSINESS_PNG   = os.path.join(SCREENSHOTS_DIR, 'itf_business_flow.png')
FLOW_APPROVAL_PNG   = os.path.join(SCREENSHOTS_DIR, 'itf_approval_flow.png')
ERD_MAIN_PNG        = os.path.join(SCREENSHOTS_DIR, 'itf_erd_main.png')

# Screenshot files (captured from browser - stored in BRAIN_DIR directly)
SS_INDEX      = os.path.join(BRAIN_DIR, 'trial_formulation_index_1775791883114.png')
SS_TAB1_NEW   = os.path.join(BRAIN_DIR, 'trial_formulation_detail_tab1_1775791896874.png')
SS_TAB2_NEW   = os.path.join(BRAIN_DIR, 'trial_formulation_detail_tab2_1775791908701.png')
SS_TAB3_NEW   = os.path.join(BRAIN_DIR, 'trial_formulation_detail_tab3_1775791928999.png')
SS_TAB1_EDIT  = os.path.join(BRAIN_DIR, 'trial_formulation_edit_tab1_1775791957690.png')
SS_TAB2_EDIT  = os.path.join(BRAIN_DIR, 'trial_formulation_edit_tab2_1775791970127.png')
SS_TAB3_EDIT  = os.path.join(BRAIN_DIR, 'trial_formulation_edit_tab3_1775791982098.png')

# ──────────────────────────────────────────────────
# STYLE CONSTANTS (sesuai standar FSD sebelumnya)
# ──────────────────────────────────────────────────
HEADER_BG       = 'D9EAD3'   # light green
BORDER_COLOR    = '000000'   # black
FONT_NAME       = 'Calibri'
FONT_SIZE_BODY  = 11
FONT_SIZE_TABLE = 9


# ──────────────────────────────────────────────────
# STEP 1: Render Mermaid diagrams via Kroki.io
# ──────────────────────────────────────────────────

def render_kroki(mermaid_code: str, output_path: str, label: str) -> bool:
    """Render Mermaid code to PNG via Kroki.io API."""
    try:
        compressed = zlib.compress(mermaid_code.strip().encode('utf-8'), 9)
        b64 = base64.urlsafe_b64encode(compressed).decode('ascii')
        url = f'https://kroki.io/mermaid/png/{b64}'
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=45) as resp:
            data = resp.read()
        with open(output_path, 'wb') as f:
            f.write(data)
        print(f'   [{label}] [OK] {len(data):,} bytes -> {os.path.basename(output_path)}')
        return True
    except Exception as e:
        print(f'   [{label}] [X] Kroki gagal: {e}')
        return False


def step1_render_diagrams(md_content: str):
    """Extract mermaid blocks and render each to PNG."""
    print('\n[STEP 1] Render diagram Mermaid via Kroki.io...')
    
    blocks = list(re.finditer(r'```mermaid\s*\n(.*?)```', md_content, flags=re.DOTALL))
    print(f'   Ditemukan {len(blocks)} blok Mermaid.')
    if len(blocks) == 0:
        print('   [WARN] Tidak ada blok mermaid ditemukan - diagram akan dilewati.')
        return

    for i, match in enumerate(blocks):
        code = match.group(1).strip()
        label = f'Diagram-{i+1}'
        
        # Identifikasi diagram berdasarkan konten
        if 'InitReq' in code and 'CreateTrialFormula' in code:
            render_kroki(code, FLOW_BUSINESS_PNG, 'Flow-Business')
        elif 'sequenceDiagram' in code and 'Creator' in code:
            render_kroki(code, FLOW_APPROVAL_PNG, 'Flow-Approval')
        elif 'erDiagram' in code and 'ITF_Header' in code:
            render_kroki(code, ERD_MAIN_PNG, 'ERD-Main')
        else:
            generic_path = os.path.join(SCREENSHOTS_DIR, f'itf_diagram_{i+1}.png')
            render_kroki(code, generic_path, label)


# ──────────────────────────────────────────────────
# STEP 2: Preprocess Markdown
# ──────────────────────────────────────────────────

def _rel(path: str) -> str:
    return os.path.relpath(path, SCRIPT_DIR).replace('\\', '/')


def step2_preprocess_markdown():
    print('\n[STEP 2] Pre-processing Markdown...')
    with open(MD_SRC, 'r', encoding='utf-8') as f:
        text = f.read()

    # Pertama, render semua diagram
    step1_render_diagrams(text)

    diagram_counter = [0]

    def replace_mermaid_block(m):
        code = m.group(1).strip()
        diagram_counter[0] += 1
        i = diagram_counter[0]

        if 'InitReq' in code and 'CreateTrialFormula' in code:
            png     = FLOW_BUSINESS_PNG
            caption = 'Business Flow Diagram – Item Trial Formulation'
        elif 'sequenceDiagram' in code and 'Creator' in code:
            png     = FLOW_APPROVAL_PNG
            caption = 'Approval Flow – Item Trial Formulation'
        elif 'erDiagram' in code and 'ITF_Header' in code:
            png     = ERD_MAIN_PNG
            caption = 'ERD – Entity Relationship Diagram Item Trial Formulation'
        else:
            generic_path = os.path.join(SCREENSHOTS_DIR, f'itf_diagram_{i}.png')
            png     = generic_path
            caption = f'Diagram {i}'

        if os.path.exists(png):
            return f'\n![{caption}]({_rel(png)})\n'
        else:
            return f'\n> *[{caption} – diagram tidak dapat di-render]*\n'

    text = re.sub(r'```mermaid\s*\n(.*?)```', replace_mermaid_block, text, flags=re.DOTALL)

    # Replace screenshot references dengan relative paths
    screenshot_mappings = {
        '../trial_formulation_index_1775791883114.png': _rel(SS_INDEX),
        '../trial_formulation_detail_tab1_1775791896874.png': _rel(SS_TAB1_NEW),
        '../trial_formulation_detail_tab2_1775791908701.png': _rel(SS_TAB2_NEW),
        '../trial_formulation_detail_tab3_1775791928999.png': _rel(SS_TAB3_NEW),
        '../trial_formulation_edit_tab1_1775791957690.png': _rel(SS_TAB1_EDIT),
        '../trial_formulation_edit_tab2_1775791970127.png': _rel(SS_TAB2_EDIT),
        '../trial_formulation_edit_tab3_1775791982098.png': _rel(SS_TAB3_EDIT),
    }

    for old_ref, new_ref_rel in screenshot_mappings.items():
        # Check existence using absolute path
        abs_path = os.path.normpath(os.path.join(SCRIPT_DIR, new_ref_rel))
        if os.path.exists(abs_path):
            text = text.replace(old_ref, new_ref_rel)
            print(f'   [OK] Screenshot resolved: {os.path.basename(abs_path)}')
        else:
            print(f'   [WARN] Screenshot tidak ditemukan: {abs_path}')

    with open(MD_TMP, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f'   [OK] Preprocessed MD -> {os.path.basename(MD_TMP)}')


# ──────────────────────────────────────────────────
# STEP 3: Run Pandoc (MD → DOCX)
# ──────────────────────────────────────────────────

def step3_run_pandoc():
    print('\n[STEP 3] Menjalankan Pandoc (MD -> DOCX)...')
    cmd = [
        'pandoc', MD_TMP,
        '-o', DOCX_OUT,
        '--from=markdown+pipe_tables',
        f'--resource-path={SCRIPT_DIR};{SCREENSHOTS_DIR}',
        '--toc',
        '--toc-depth=3',
    ]
    if REF_DOCX and os.path.exists(REF_DOCX):
        cmd += [f'--reference-doc={REF_DOCX}']
        print(f'   [OK] Using reference docx: {os.path.basename(REF_DOCX)}')

    result = subprocess.run(cmd, capture_output=True, text=True, cwd=SCRIPT_DIR)
    if result.returncode != 0:
        print('   pandoc GAGAL!\n', result.stderr[:800])
        raise RuntimeError('pandoc execution error')
    print(f'   [OK] DOCX berhasil: {os.path.basename(DOCX_OUT)}')


# ──────────────────────────────────────────────────
# STEP 4: Post-process DOCX
# ──────────────────────────────────────────────────

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = f'w:{edge}'
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for k, v in kwargs[edge].items():
                element.set(qn(f'w:{k}'), str(v))


def set_cell_bg(cell, color: str):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def apply_font(run, size_pt: int, bold=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    el = rPr.find(qn('w:rFonts'))
    if el is None:
        el = OxmlElement('w:rFonts')
        rPr.insert(0, el)
    el.set(qn('w:ascii'), FONT_NAME)
    el.set(qn('w:hAnsi'), FONT_NAME)
    el.set(qn('w:cs'), FONT_NAME)
    if bold is not None:
        run.bold = bold


def step4_postprocess_docx():
    print('\n[STEP 4] Post-processing DOCX (Fonts, Table Borders, Image Scale)...')
    doc = Document(DOCX_OUT)

    # Global font
    try:
        doc.styles['Normal'].font.name = FONT_NAME
        doc.styles['Normal'].font.size = Pt(FONT_SIZE_BODY)
    except Exception:
        pass

    # Apply font to paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            apply_font(run, FONT_SIZE_BODY)

    border_spec = {"sz": "8", "val": "single", "color": BORDER_COLOR}
    bdr_all = dict(
        top=border_spec, bottom=border_spec,
        start=border_spec, end=border_spec,
        insideH=border_spec, insideV=border_spec
    )

    # Style tables
    for table in doc.tables:
        try:
            table.style = 'Table Grid'
        except Exception:
            pass

        for row_idx, row in enumerate(table.rows):
            is_header = (row_idx == 0)
            for cell in row.cells:
                set_cell_border(cell, **bdr_all)
                if is_header:
                    set_cell_bg(cell, HEADER_BG)
                for para in cell.paragraphs:
                    for run in para.runs:
                        apply_font(run, FONT_SIZE_TABLE, bold=True if is_header else None)

    # Scale images to max width 15cm
    MAX_WIDTH = Cm(15)
    for para in doc.paragraphs:
        for run in para.runs:
            drawings = run._r.findall(
                './/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'
            )
            for drawing in drawings:
                extent = drawing.find(
                    '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'
                )
                if extent is not None:
                    cx = int(extent.get('cx', 0))
                    if cx > MAX_WIDTH.emu:
                        ratio = MAX_WIDTH.emu / cx
                        cy = int(int(extent.get('cy', 0)) * ratio)
                        extent.set('cx', str(MAX_WIDTH.emu))
                        extent.set('cy', str(cy))

    doc.save(DOCX_OUT)
    print(f'   [OK] DOCX final tersimpan: {os.path.basename(DOCX_OUT)}')


# ──────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────

if __name__ == '__main__':
    print('=' * 65)
    print('  BUILD START: FSD Item Trial Formulation v1.0')
    print('  Migration dari KN2017_FORMULATION ke IDC System')
    print('=' * 65)

    # Verify source markdown exists
    if not os.path.exists(MD_SRC):
        print(f'\n  ERROR: Source MD tidak ditemukan: {MD_SRC}')
        print('  Pastikan FSD_Item_Trial_Formulation_v1.0.md ada di folder yang sama.')
        exit(1)

    print(f'\n  Source MD : {os.path.basename(MD_SRC)}')
    print(f'  Output    : {os.path.basename(DOCX_OUT)}')
    print(f'  Screenshots: {SCREENSHOTS_DIR}')
    if REF_DOCX:
        print(f'  Reference : {os.path.basename(REF_DOCX)}')
    else:
        print(f'  Reference : (tidak ditemukan - menggunakan default pandoc)')

    step2_preprocess_markdown()
    step3_run_pandoc()
    step4_postprocess_docx()

    # Cleanup temp file
    if os.path.exists(MD_TMP):
        os.remove(MD_TMP)
        print(f'\n   [OK] Temp file dihapus: {os.path.basename(MD_TMP)}')

    print('\n' + '=' * 65)
    print(f'  SELESAI: {DOCX_OUT}')
    print('=' * 65)
