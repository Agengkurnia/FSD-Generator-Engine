"""
build_formula_fsd.py
====================
Pipeline untuk membangun 6 FSD Formulation Modules ke format DOCX:
  1. Premix Formula
  2. DryBlend Formula
  3. Liquid Formula
  4. Baking Formula
  5. Drying Formula
  6. Packaging Formula

Run:
    py build_formula_fsd.py
    py build_formula_fsd.py --module premix
"""

import base64, os, re, subprocess, shutil, argparse, urllib.request, zlib
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# PATHS
# ─────────────────────────────────────────────
SCRIPT_DIR   = os.path.dirname(os.path.abspath(__file__))
FSD_DIR      = SCRIPT_DIR                  # same folder as this script
SCREENSHOTS  = os.path.join(SCRIPT_DIR, 'screenshots')
REF_DOCX     = os.path.join(os.path.dirname(SCRIPT_DIR), 'reference.docx')

os.makedirs(SCREENSHOTS, exist_ok=True)

# ─────────────────────────────────────────────
# STYLING
# ─────────────────────────────────────────────
HEADER_BG       = 'BDD7EE'   # light blue (sesuai standar IDC System)
BORDER_COLOR    = '000000'
FONT_NAME       = 'Calibri'
FONT_SIZE_BODY  = 11
FONT_SIZE_TABLE = 9

# ─────────────────────────────────────────────
# MODULE DEFINITIONS
# ─────────────────────────────────────────────
MODULES = {
    'premix':    'FSD_PremixFormula_v1.0',
    'dryblend':  'FSD_DryBlendFormula_v1.0',
    'liquid':    'FSD_LiquidFormula_v1.0',
    'baking':    'FSD_BakingFormula_v1.0',
    'drying':    'FSD_DryingFormula_v1.0',
    'packaging': 'FSD_PackagingFormula_v1.0',
    'combined':  'FSD_FormulationModules_v1.0',
}

# ─────────────────────────────────────────────
# MERMAID RENDERER (via kroki.io)
# ─────────────────────────────────────────────
def render_mermaid_png(mermaid_code, output_path, label):
    compressed = zlib.compress(mermaid_code.strip().encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('ascii')
    url = f'https://kroki.io/mermaid/png/{encoded}'
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = resp.read()
        with open(output_path, 'wb') as f:
            f.write(data)
        print(f'   [{label}] OK {len(data):,} bytes -> {os.path.basename(output_path)}')
        return output_path
    except Exception as e:
        print(f'   [{label}] FAIL: {e}')
        return None

# ─────────────────────────────────────────────
# MARKDOWN PREPROCESSOR
# ─────────────────────────────────────────────
def preprocess_markdown(md_src, md_tmp):
    with open(md_src, 'r', encoding='utf-8') as f:
        text = f.read()

    counter = [0]
    def replace_mermaid(m):
        counter[0] += 1
        code = m.group(1).strip()
        label = f'diagram_{counter[0]}'
        png_path = os.path.join(SCREENSHOTS, f'{label}.png')
        result = render_mermaid_png(code, png_path, label)
        if result and os.path.exists(result):
            rel = os.path.relpath(result, os.path.dirname(md_tmp)).replace('\\', '/')
            return f'\n![Diagram]({rel})\n'
        return f'\n*(Diagram tidak tersedia)*\n'

    text = re.sub(r'```mermaid\s*\n(.*?)```', replace_mermaid, text, flags=re.DOTALL)

    with open(md_tmp, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f'   Markdown preprocessed: {os.path.basename(md_tmp)}')

# ─────────────────────────────────────────────
# PANDOC RUNNER
# ─────────────────────────────────────────────
def run_pandoc(md_tmp, docx_out):
    cmd = [
        'pandoc', md_tmp,
        '-o', docx_out,
        '--from=markdown+pipe_tables',
        f'--resource-path={FSD_DIR};{SCREENSHOTS}',
        '--toc',
        '--toc-depth=3',
    ]
    if os.path.exists(REF_DOCX):
        cmd += [f'--reference-doc={REF_DOCX}']
    result = subprocess.run(cmd, capture_output=True, text=True, cwd=FSD_DIR)
    if result.returncode != 0:
        print(f'   pandoc error: {result.stderr[:300]}')
        raise RuntimeError('pandoc gagal!')
    print(f'   DOCX dibuat: {os.path.basename(docx_out)}')

# ─────────────────────────────────────────────
# DOCX POST-PROCESSOR
# ─────────────────────────────────────────────
def set_cell_border(cell, color='000000'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        tag = f'w:{edge}'
        el = tcBorders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tcBorders.append(el)
        el.set(qn('w:sz'), '4')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:color'), color)

def set_cell_bg(cell, color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)

def apply_font(run, size_pt, bold=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    el = rPr.find(qn('w:rFonts'))
    if el is None:
        el = OxmlElement('w:rFonts')
        rPr.insert(0, el)
    el.set(qn('w:ascii'), FONT_NAME)
    el.set(qn('w:hAnsi'), FONT_NAME)
    el.set(qn('w:cs'), FONT_NAME)
    if bold is not None:
        run.bold = bold

def postprocess_docx(docx_out):
    doc = Document(docx_out)
    normal = doc.styles['Normal']
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE_BODY)

    for para in doc.paragraphs:
        for run in para.runs:
            apply_font(run, FONT_SIZE_BODY)

    for table in doc.tables:
        try:
            table.style = 'Table Grid'
        except Exception:
            pass
        for row_idx, row in enumerate(table.rows):
            is_header = (row_idx == 0)
            for cell in row.cells:
                set_cell_border(cell, BORDER_COLOR)
                if is_header:
                    set_cell_bg(cell, HEADER_BG)
                for para in cell.paragraphs:
                    for run in para.runs:
                        apply_font(run, FONT_SIZE_TABLE, bold=True if is_header else None)

    # Limit image width to 14cm
    MAX_WIDTH = Cm(14)
    for para in doc.paragraphs:
        for run in para.runs:
            for drawing in run._r.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
                extent = drawing.find('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                if extent is not None:
                    cx = int(extent.get('cx', 0))
                    if cx > MAX_WIDTH.emu:
                        ratio = MAX_WIDTH.emu / cx
                        cy = int(int(extent.get('cy', 0)) * ratio)
                        extent.set('cx', str(MAX_WIDTH.emu))
                        extent.set('cy', str(cy))

    doc.save(docx_out)
    print(f'   DOCX post-processed: {os.path.basename(docx_out)}')

# ─────────────────────────────────────────────
# BUILD ONE MODULE
# ─────────────────────────────────────────────
def build_module(key, basename):
    print(f'\n{"="*55}')
    print(f'  Building: {basename}')
    print(f'{"="*55}')

    md_src  = os.path.join(FSD_DIR, f'{basename}.md')
    md_tmp  = os.path.join(FSD_DIR, f'_tmp_{basename}.md')
    docx_out = os.path.join(FSD_DIR, f'{basename}.docx')

    if not os.path.exists(md_src):
        print(f'   SKIP: {md_src} tidak ditemukan')
        return False

    try:
        preprocess_markdown(md_src, md_tmp)
        run_pandoc(md_tmp, docx_out)
        postprocess_docx(docx_out)
        return True
    except Exception as e:
        print(f'   ERROR: {e}')
        return False
    finally:
        if os.path.exists(md_tmp):
            os.remove(md_tmp)

# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Build FSD Formulation Modules ke DOCX')
    parser.add_argument('--module', choices=list(MODULES.keys()) + ['all'], default='all',
                        help='Module yang akan dibangun (default: all)')
    args = parser.parse_args()

    print('\n' + '='*55)
    print('  FSD Formulation Modules – DOCX Builder')
    print('='*55)

    target = {args.module: MODULES[args.module]} if args.module != 'all' else MODULES

    results = {}
    for key, basename in target.items():
        results[key] = build_module(key, basename)

    print('\n' + '='*55)
    print('  SUMMARY')
    print('='*55)
    for key, ok in results.items():
        status = 'OK' if ok else 'FAIL'
        print(f'  [{status}] {MODULES.get(key, key)}')
    print('='*55)
    print(f'\n  Output folder: {FSD_DIR}')
